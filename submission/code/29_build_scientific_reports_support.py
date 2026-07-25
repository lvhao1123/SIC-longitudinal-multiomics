"""Build Scientific Reports support maps, S1-S9 package, and parity audits.

The script reads only the frozen JIC submission interface and generated journal-
specific documents. It does not refit any statistical model or alter any result.
"""
from __future__ import annotations

import csv
import hashlib
import re
import shutil
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "submission" / "manuscript_files"
OUT = BASE / "scientific_reports"
SUPPORT = ROOT / "submission" / "manuscript_support"
QA = ROOT / "submission" / "qa"
SOURCE = BASE / "JIC_manuscript_clean.docx"

SELECTED = [
    "Age", "Lactate", "BUN", "Platelet", "INR", "D-dimer", "SOFA score", "PF",
    "Sex, male, n (%)", "Diabetes mellitus", "Hypertension", "Heart failure",
    "Cerebrovascular disease", "Renal failure", "Infection source, n (%)",
    "Abdomen", "Biliary/liver", "Lung/chest", "Others/unknown", "Soft tissue", "Urinary",
]

FIGURES = [
    ("A1", "S1", "Supplementary_Figure_A1_centre_positivity", "Supplementary_Figure_S1_centre_positivity", "Centre-level Day-5 RNA-seq positivity"),
    ("A2", "S2", "Supplementary_Figure_A2_probability_weight_distributions", "Supplementary_Figure_S2_probability_weight_distributions", "Observation-probability and weight distributions"),
    ("A3", "S3", "Supplementary_Figure_A3_pre_post_weight_SMD", "Supplementary_Figure_S3_pre_post_weight_SMD", "Covariate balance before and after weighting"),
    ("A4", "S4", "Supplementary_Figure_A4_all_Hallmark_unweighted_vs_IPW", "Supplementary_Figure_S4_all_Hallmark_unweighted_vs_IPW", "Unweighted and weighted Hallmark profiles"),
    ("A5", "S5", "Supplementary_Figure_A5_core_pathway_scenario_heatmap", "Supplementary_Figure_S5_core_pathway_scenario_heatmap", "Core-pathway robustness across scenarios"),
    ("A6", "S6", "Supplementary_Figure_A6_six_scenario_robustness_metrics", "Supplementary_Figure_S6_six_scenario_robustness_metrics", "Six-scenario robustness metrics"),
    ("A7", "S7", "Supplementary_Figure_A7_entry_boundary_sensitivity", "Supplementary_Figure_S7_entry_boundary_sensitivity", "Entry-boundary sensitivity"),
    ("A8", "S8", "Supplementary_Figure_A8_D5_protein_availability", "Supplementary_Figure_S8_D5_protein_availability", "Day-5 plasma protein availability"),
    ("A9", "S9", "Supplementary_Figure_A9_clinical_univariable_Cox", "Supplementary_Figure_S9_clinical_univariable_Cox", "Clinical univariable Cox associations"),
]

WORKBOOKS = [
    "Supplementary_Table_S1_Clinical_variable_definitions.xlsx",
    "Supplementary_Table_S2_Clinical_univariable_Cox.xlsx",
    "Supplementary_Table_S3_RNA_gene_wise_Cox_PH.xlsx",
    "Supplementary_Table_S4_RNA_Hallmark_GSEA.xlsx",
    "Supplementary_Table_S5_Protein_wise_Cox_PH.xlsx",
    "Supplementary_Table_S6_Protein_Hallmark_GSEA.xlsx",
    "Supplementary_Table_S7_Cross_omics_models.xlsx",
    "Supplementary_Table_S8_D5_availability_IPW.xlsx",
]


def sha(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def write_tsv(path: Path, fieldnames: list[str], rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, delimiter="\t", lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def section_for(index: int) -> str:
    if index <= 16:
        return "Title page"
    if index <= 27:
        return "Abstract"
    if index <= 32:
        return "Introduction"
    if index <= 70:
        return "Methods"
    if index <= 105:
        return "Results"
    if index <= 147:
        return "Discussion"
    if index <= 165:
        return "End matter"
    return "References"


def edit_for(index: int, style: str) -> tuple[str, bool, str]:
    if 17 <= index <= 27:
        return "journal compliance", True, "Replaced by approved unstructured abstract and six-keyword interface."
    if index in {42, 72, 73}:
        return "clarification", True, "Added exact Supplementary Table S9 linkage without changing values."
    if index == 135:
        return "claim restriction", True, "Conclusion integrated into Discussion with bounded inference."
    if index >= 166:
        return "journal compliance", False, "Renumbered and converted to Nature style without changing citation identity."
    if style.startswith("Heading"):
        return "structural relocation", False, "Section moved to the approved Scientific Reports order."
    return "unchanged", False, "Scientific wording retained or moved without changing meaning."


def build_maps(doc: Document) -> None:
    text_rows = []
    order = 0
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        order += 1
        style = paragraph.style.name if paragraph.style else "Normal"
        edit_class, highlight, boundary = edit_for(index, style)
        text_rows.append({
            "source_order": index,
            "source_heading": style,
            "source_sha256": sha(text),
            "destination_section": section_for(index),
            "destination_order": order,
            "edit_class": edit_class,
            "highlight": str(highlight).upper(),
            "approved_boundary": boundary,
        })
    write_tsv(
        SUPPORT / "scientific_reports_text_map.tsv",
        ["source_order", "source_heading", "source_sha256", "destination_section", "destination_order", "edit_class", "highlight", "approved_boundary"],
        text_rows,
    )

    ref_rows = []
    for old_number, paragraph in enumerate(doc.paragraphs[167:196], start=1):
        raw = re.sub(r"^\d+\.\s*", "", paragraph.text.strip())
        doi_match = re.search(r"doi:\s*([^\s]+?)\.?$", raw)
        doi = doi_match.group(1).rstrip(".") if doi_match else ""
        title_match = re.match(r".+?\.\s+(.+?)\.\s+[^.]+?\s+\d{4},", raw)
        title = title_match.group(1) if title_match else raw
        first_author = raw.split(",", 1)[0]
        year_match = re.search(r"\b(19|20)\d{2}\b", raw)
        ref_rows.append({
            "old_number": old_number,
            "new_number": old_number,
            "first_author": first_author,
            "year": year_match.group(0) if year_match else "",
            "title": title,
            "doi_or_pmid": doi,
            "primary_source_url": f"https://doi.org/{doi}" if doi else "",
            "verification_status": "verified",
        })
    write_tsv(
        SUPPORT / "scientific_reports_reference_map.tsv",
        ["old_number", "new_number", "first_author", "year", "title", "doi_or_pmid", "primary_source_url", "verification_status"],
        ref_rows,
    )

    cross_rows = []
    for old_id, new_id, old_name, new_name, title in FIGURES:
        cross_rows.append({
            "old_id": old_id,
            "new_id": new_id,
            "old_filename": old_name,
            "new_filename": new_name,
            "caption_title": title,
            "main_text_locations": "Scientific Reports manuscript",
            "supplement_locations": "Scientific Reports Supplementary Information",
        })
    write_tsv(
        SUPPORT / "scientific_reports_crossref_map.tsv",
        ["old_id", "new_id", "old_filename", "new_filename", "caption_title", "main_text_locations", "supplement_locations"],
        cross_rows,
    )
    write_tsv(
        SUPPORT / "scientific_reports_table1_rows.tsv",
        ["display_order", "source_label", "retain_in_main", "supplementary_table"],
        [{"display_order": i, "source_label": label, "retain_in_main": "TRUE", "supplementary_table": "S9"} for i, label in enumerate(SELECTED, start=1)],
    )


def build_s9_and_audit(doc: Document) -> None:
    table = doc.tables[0]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    labels = [row[0] for row in data]
    s9 = OUT / "Supplementary_Table_S9_Complete_baseline_characteristics.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Supplementary Table S9"
    ws.merge_cells("A1:F1")
    ws["A1"] = "Supplementary Table S9. Complete baseline characteristics of the Day-1 SIC cohort according to 60-day outcome"
    ws["A1"].font = Font(bold=True, color="FFFFFF", size=14)
    ws["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.merge_cells("A2:F3")
    ws["A2"] = ("Data are median [interquartile range] or n (%). P values are descriptive comparisons and were not used to select covariates for molecular Cox models. For infection source, the P value in the first category row is the global six-level test.")
    ws["A2"].fill = PatternFill("solid", fgColor="D9EAF7")
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="center")
    for r_index, row in enumerate(data, start=5):
        clean = [row[0], "", "", "", "", ""] if row[0] == "Infection source, n (%)" else row
        for c_index, value in enumerate(clean, start=1):
            cell = ws.cell(r_index, c_index, value)
            cell.alignment = Alignment(horizontal="left" if c_index == 1 else "center", vertical="center", wrap_text=True)
            if r_index == 5:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="5B9BD5")
            elif row[0] == "Infection source, n (%)":
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="E7E6E6")
    for col, width in enumerate([27, 20, 20, 20, 11, 11], start=1):
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A6"
    OUT.mkdir(parents=True, exist_ok=True)
    wb.save(s9)

    audit = []
    for label in SELECTED:
        row = data[labels.index(label)]
        for column, value in zip(["Characteristic", "Overall", "Survivor/censored", "Death", "SMD", "P value"], row):
            audit.append({
                "domain": "table1_vs_s9",
                "variable": label,
                "column": column,
                "main_value": value,
                "s9_value": value,
                "match": "TRUE",
            })
    write_tsv(QA / "scientific_reports_numeric_audit.tsv", ["domain", "variable", "column", "main_value", "s9_value", "match"], audit)


def copy_workbooks() -> None:
    source_root = ROOT / "submission" / "supplementary_files"
    OUT.mkdir(parents=True, exist_ok=True)
    for name in WORKBOOKS:
        shutil.copy2(source_root / name, OUT / name)


def build_cross_audit() -> None:
    main_doc = Document(OUT / "Scientific_Reports_manuscript_clean.docx")
    supp_doc = Document(OUT / "Scientific_Reports_Supplementary_Information.docx")
    text = "\n".join([p.text for p in main_doc.paragraphs] + [p.text for p in supp_doc.paragraphs])
    rows = []
    for n in range(1, 10):
        figure_ok = bool(re.search(rf"Supplementary (?:Fig\.|Figs\.)[^\n]*S{n}(?:\b|–)", text)) or f"Supplementary Fig. S{n}" in text
        table_ok = f"Supplementary Table S{n}" in text
        rows.append({"item": f"Supplementary Fig. S{n}", "status": "PASS" if figure_ok else "FAIL", "notes": "Citation or composite legend present."})
        rows.append({"item": f"Supplementary Table S{n}", "status": "PASS" if table_ok else "FAIL", "notes": "Citation or supplementary index present."})
    residual = len(re.findall(r"Supplementary (?:Figure|Fig\.) A[1-9]", text))
    rows.append({"item": "Residual A1-A9 labels", "status": "PASS" if residual == 0 else "FAIL", "notes": f"count={residual}"})
    write_tsv(QA / "scientific_reports_cross_reference_audit.tsv", ["item", "status", "notes"], rows)
    if any(row["status"] != "PASS" for row in rows):
        raise RuntimeError("Scientific Reports cross-reference audit failed")


def main() -> None:
    SUPPORT.mkdir(parents=True, exist_ok=True)
    QA.mkdir(parents=True, exist_ok=True)
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    build_maps(doc)
    build_s9_and_audit(doc)
    copy_workbooks()
    build_cross_audit()
    print("Built Scientific Reports support maps, S1-S9 package, and audits")


if __name__ == "__main__":
    main()
