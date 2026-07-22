"""Build clean and commented Critical Care manuscript revisions.

This production script edits language and manuscript interfaces only. It does
not read participant-level data or refit any statistical model.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


REVIEW_AUTHOR = "Critical Care pre-submission review"
REVIEW_INITIALS = "CC"


def insert_paragraph_after(paragraph, text: str, style: str = "Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def replace_paragraph(doc, old: str, new: str, comment: str | None, annotated: bool):
    matches = [p for p in doc.paragraphs if p.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one paragraph match, found {len(matches)}: {old[:80]}")
    paragraph = matches[0]
    paragraph.clear()
    run = paragraph.add_run(new)
    if annotated and comment:
        doc.add_comment(
            run,
            text=comment,
            author=REVIEW_AUTHOR,
            initials=REVIEW_INITIALS,
        )
    return paragraph


def comment_paragraph(doc, paragraph, text: str, annotated: bool):
    if not annotated:
        return
    runs = [run for run in paragraph.runs if run.text]
    if runs:
        doc.add_comment(
            runs,
            text=text,
            author=REVIEW_AUTHOR,
            initials=REVIEW_INITIALS,
        )


def replace_embedded_figure(doc, caption_prefix: str, image_path: Path, width_inches: float):
    """Replace the nearest drawing before a uniquely identified figure caption."""
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    captions = [p for p in doc.paragraphs if p.text.startswith(caption_prefix)]
    if len(captions) != 1:
        raise RuntimeError(
            f"Expected one caption beginning {caption_prefix!r}; found {len(captions)}"
        )
    node = captions[0]._p.getprevious()
    while node is not None:
        paragraph = Paragraph(node, captions[0]._parent)
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
            return
        if paragraph.text.startswith("Supplementary Figure"):
            break
        node = node.getprevious()
    raise RuntimeError(f"No embedded drawing found before {caption_prefix!r}")


def set_run_font(run, name: str = "Arial", size_pt: float = 9, bold: bool | None = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float):
    width_dxa = str(round(width_inches * 1440))
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), width_dxa)
    tc_w.set(qn("w:type"), "dxa")


def bind_main_figure_block(doc, figure_number: int, short_title: str, image_path: Path, width_inches: float):
    """Replace the main image and retain its full legend; the display title is embedded in the image."""
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    display_matches = [p for p in doc.paragraphs if p.text == f"Figure {figure_number}"]
    legend_matches = [
        p for p in doc.paragraphs
        if p.text.startswith(f"Figure {figure_number} |") and len(p.text) > 80
    ]
    if len(display_matches) != 1 or len(legend_matches) != 1:
        raise RuntimeError(
            f"Expected one display heading and one full legend for Figure {figure_number}; "
            f"found {len(display_matches)} and {len(legend_matches)}"
        )
    heading = display_matches[0]
    node = heading._p.getnext()
    if node is None or not node.xpath(".//w:drawing"):
        raise RuntimeError(f"Figure {figure_number} drawing does not follow its display heading")
    drawing = Paragraph(node, heading._parent)

    drawing.clear()
    drawing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    drawing.paragraph_format.keep_with_next = True
    drawing.paragraph_format.page_break_before = True
    drawing.paragraph_format.space_before = Pt(0)
    drawing.paragraph_format.space_after = Pt(0)
    drawing.add_run().add_picture(str(image_path), width=Inches(width_inches))
    remove_paragraph(heading)

    legend = legend_matches[0]
    original = legend.text
    old_title, separator, remainder = original.partition(". ")
    if not separator:
        raise RuntimeError(f"Figure {figure_number} legend has no title separator")
    legend.clear()
    title_run = legend.add_run(f"Figure {figure_number} | {short_title}.")
    set_run_font(title_run, size_pt=9, bold=True)
    body_run = legend.add_run(f" {remainder}")
    set_run_font(body_run, size_pt=9, bold=False)
    legend.style = "Figure note"
    legend.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    legend.paragraph_format.line_spacing = 1.05
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(8)
    legend.paragraph_format.keep_with_next = False
    drawing._p.addnext(legend._p)


def remove_paragraph(paragraph):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def standardise_table1_labels(doc, annotated: bool):
    """Standardise clinical labels without changing any Table 1 result value."""
    if len(doc.tables) != 1:
        raise RuntimeError(f"Expected one manuscript table; found {len(doc.tables)}")
    table = doc.tables[0]
    if len(table.rows) != 46 or len(table.columns) != 6:
        raise RuntimeError(
            f"Unexpected Table 1 dimensions: {len(table.rows)} x {len(table.columns)}"
        )

    labels = {
        "Body mass index, kg/m^2": "BMI, kg/m²",
        "Maximum heart rate, beats/min": "HRmax, beats/min",
        "Maximum mean arterial pressure, mmHg": "MAPmax, mmHg",
        "Maximum systolic arterial pressure, mmHg": "SBPmax, mmHg",
        "Maximum respiratory rate, breaths/min": "RRmax, breaths/min",
        "Maximum temperature, degrees C": "Tmax, °C",
        "Potassium, mmol/L": "K, mmol/L",
        "Sodium, mmol/L": "Na, mmol/L",
        "Chloride, mmol/L": "Cl, mmol/L",
        "Blood urea nitrogen, mg/dL": "BUN, mg/dL",
        "Creatinine, micromol/L": "Creatinine, µmol/L",
        "Bilirubin, micromol/L": "Bilirubin, µmol/L",
        "C-reactive protein, mg/L": "CRP, mg/L",
        "Procalcitonin, ng/mL": "PCT, ng/mL",
        "White blood cell count, 10^9/L": "WBC count, ×10⁹/L",
        "Platelet count, 10^9/L": "Platelet count, ×10⁹/L",
        "International normalized ratio, ratio": "INR",
        "Activated partial thromboplastin time, seconds": "aPTT, s",
        "D-dimer, microg/mL": "D-dimer, µg/mL",
        "PaO2/FiO2 ratio, mmHg": "PaO₂/FiO₂ ratio, mmHg",
        "Male sex: 1 vs 0": "Male sex",
        "Diabetes mellitus: 1 vs 0": "Diabetes mellitus",
        "Hypertension: 1 vs 0": "Hypertension",
        "Myocardial infarction: 1 vs 0": "Myocardial infarction",
        "Heart failure: 1 vs 0": "Heart failure",
        "Cerebrovascular disease: 1 vs 0": "Cerebrovascular disease",
        "Dementia: 1 vs 0": "Dementia",
        "Chronic obstructive pulmonary disease: 1 vs 0": "COPD",
        "Paralysis: 1 vs 0": "Paralysis",
        "Renal failure: 1 vs 0": "Renal failure",
        "Infection source: Abdomen": "Abdomen",
        "  Biliary/Liver": "Biliary/liver",
        "  Lung/Chest": "Lung/chest",
        "  Others/Unknown": "Others/unknown",
        "  SoftTissue": "Soft tissue",
        "  Urinary": "Urinary",
    }
    removed_labels = {"Continuous variables", "Binary variables"}
    result_cells_before = [
        [cell.text for cell in row.cells[1:]]
        for row in table.rows
        if row.cells[0].text not in removed_labels
    ]
    observed = {row.cells[0].text for row in table.rows}
    missing = sorted(set(labels) - observed)
    if missing:
        raise RuntimeError(f"Table 1 labels not found: {missing}")

    for row in list(table.rows)[::-1]:
        if row.cells[0].text in removed_labels:
            table._tbl.remove(row._tr)

    for row in table.rows:
        old = row.cells[0].text
        if old in labels:
            paragraphs = row.cells[0].paragraphs
            runs = [run for paragraph in paragraphs for run in paragraph.runs]
            if not runs:
                raise RuntimeError(f"Table 1 label has no editable run: {old}")
            runs[0].text = labels[old]
            for run in runs[1:]:
                run.text = ""

    if len(table.rows) != 44:
        raise RuntimeError(f"Unexpected Table 1 row count after removing section rows: {len(table.rows)}")

    result_cells_after = [
        [cell.text for cell in row.cells[1:]]
        for row in table.rows
    ]
    if result_cells_before != result_cells_after:
        raise RuntimeError("Table 1 result cells changed while standardising labels")

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_widths = [1.80, 1.15, 1.35, 1.15, .50, .55]
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size_pt=8, bold=True if row_index == 0 else None)

    old_note = (
        "Data are median [interquartile range] or n (%). P values describe baseline group comparisons and were "
        "not used to select covariates for the molecular Cox models. SMD, standardised mean difference; SIC, "
        "sepsis-induced coagulopathy; SOFA, Sequential Organ Failure Assessment."
    )
    new_note = (
        "Data are median [interquartile range] or n (%). P values describe baseline group comparisons and were "
        "not used to select covariates for the molecular Cox models. aPTT, activated partial thromboplastin "
        "time; BMI, body mass index; BUN, blood urea nitrogen; Cl, chloride; COPD, chronic obstructive pulmonary "
        "disease; CRP, C-reactive protein; HRmax, maximum heart rate; INR, international normalised ratio; "
        "K, potassium; MAPmax, maximum mean arterial pressure; Na, sodium; PCT, procalcitonin; RRmax, maximum "
        "respiratory rate; SBPmax, maximum systolic arterial pressure; SIC, sepsis-induced coagulopathy; SMD, "
        "standardised mean difference; SOFA, Sequential Organ Failure Assessment; Tmax, maximum temperature; "
        "WBC, white blood cell."
    )
    note = replace_paragraph(
        doc,
        old_note,
        new_note,
        "Standardised Table 1 labels, clinical abbreviations and units. No numerical value, P value or SMD was changed.",
        annotated,
    )
    return note


def apply_revision(
    source: Path,
    output: Path,
    main_figures: dict[int, Path],
    a7: Path,
    a9: Path,
    annotated: bool,
):
    doc = Document(source)

    replacements = [
        (
            "Landmark-specific mortality-associated whole-blood transcriptomic and plasma proteomic programmes in Day-1-defined sepsis-induced coagulopathy: a secondary analysis of a multicentre longitudinal cohort",
            "Landmark-specific whole-blood transcriptomic and plasma proteomic programmes associated with mortality in sepsis-induced coagulopathy: a multicentre longitudinal cohort study",
            "The title was shortened and made more readable while retaining the landmark-specific, compartment-specific and observational design. Consider adding ‘secondary analysis’ at submission only if required by the journal or editor.",
        ),
        (
            "Sepsis-induced coagulopathy (SIC) arises within a dynamic thromboinflammatory network, but mortality-associated molecular rankings may differ across serial sampling landmarks and between whole-blood and plasma compartments. We defined landmark-specific whole-blood transcriptomic and plasma proteomic mortality programmes, cross-omic relationships and selection associated with later sampling.",
            "Sepsis-induced coagulopathy (SIC) arises within a dynamic thromboinflammatory network, but the molecular programmes associated with mortality may differ across sampling landmarks and between whole-blood and plasma compartments. We defined landmark-specific transcriptomic and proteomic mortality-associated programmes, examined cross-omic relationships and assessed selection into later sampling.",
            "Tightened the study gap and replaced the ambiguous phrase ‘selection associated with later sampling’ with the operational aim of assessing selection into later sampling.",
        ),
        (
            "We analysed 504 patients meeting SIC criteria on Day 1 in the multicentre Chinese Multi-omics Advances in Sepsis cohort. Whole-blood RNA-seq and plasma protein measurements obtained on Days 1, 3 and 5 entered separate centre-stratified Cox models for 60-day mortality; later models used prespecified delayed entry so that each estimate applied only to patients eligible at the corresponding entry boundary. Cox Wald rankings underwent Hallmark gene-set enrichment. Analysis-prespecified pathway scores were evaluated in adjusted contemporaneous, forward and reverse-order cross-omic models. Day-5 RNA-seq observation was examined in a positivity-supported inverse-observation-probability-weighting sensitivity analysis.",
            "We analysed 504 patients who met SIC criteria on Day 1 in the multicentre Chinese Multi-omics Advances in Sepsis cohort. Whole-blood RNA-seq and plasma protein measurements obtained on Days 1, 3 and 5 were analysed in separate centre-stratified Cox models for 60-day mortality; the later models used prespecified delayed entry. Cox Wald rankings were subjected to Hallmark gene-set enrichment. Analysis-prespecified pathway scores were evaluated in adjusted contemporaneous, forward and reverse-order cross-omic models. Day-5 RNA-seq observation was examined in a positivity-supported inverse-observation-probability-weighting sensitivity analysis.",
            "Reduced syntactic load without changing the model, estimand or analysis hierarchy.",
        ),
        (
            "The risk-valid RNA-seq datasets included 504 patients/84 deaths, 420/67 and 320/53 on Days 1, 3 and 5; the corresponding plasma proteomic datasets included 168/27, 147/18 and 114/14. Early transcriptomic mortality rankings were positively enriched for heme metabolism, hypoxia, TNF-α/NF-κB, coagulation and reactive-oxygen-species programmes. By Day 5, MYC targets V1 (normalised enrichment score [NES] −3.13), oxidative phosphorylation (−2.44), unfolded-protein response (−2.04), E2F targets (−1.89) and DNA repair (−1.62) were negatively enriched, while heme, hypoxia and inflammatory programmes remained positive (all Benjamini–Hochberg false-discovery rate [FDR] <0.05). At the later landmarks, plasma proteomic mortality rankings were dominated by extracellular-matrix/tissue-remodelling enrichment, which remained significant on Day 5 (NES 1.94; FDR 5.67×10⁻⁶). The Day-1 transcriptomic IFN-α score was associated with the Day-3 protein IFN-α score (standardised β 0.343, 95% confidence interval 0.188–0.498; FDR 4.21×10⁻⁴). Day-3 transcriptomic IFN-α and IFN-γ scores were associated with corresponding Day-5 protein scores (FDR 0.0074 and 0.0235); no reverse-order model was FDR significant. Weighted and unweighted Day-5 Hallmark profiles were concordant (NES Spearman ρ=0.980), although positivity and residual balance remained limited.",
            "The risk-valid RNA-seq datasets comprised 504 patients with 84 deaths on Day 1, 420 with 67 deaths on Day 3 and 320 with 53 deaths on Day 5; the corresponding plasma proteomic datasets comprised 168/27, 147/18 and 114/14 patients/deaths. Early transcriptomic mortality-associated rankings were positively enriched for heme metabolism, hypoxia, TNF-α signalling via NF-κB, coagulation and reactive oxygen species pathways. By Day 5, MYC targets V1 (normalised enrichment score [NES] −3.13), oxidative phosphorylation (−2.44), unfolded protein response (−2.04), E2F targets (−1.89) and DNA repair (−1.62) were negatively enriched, whereas heme, hypoxia and inflammatory programmes remained positively enriched (all Benjamini–Hochberg false-discovery rate [FDR] <0.05). At the later landmarks, plasma proteomic mortality-associated rankings were dominated by extracellular-matrix/tissue-remodelling enrichment, which remained significant on Day 5 (NES 1.94; FDR 5.67×10⁻⁶). The Day-1 transcriptomic IFN-α score was associated with the Day-3 protein IFN-α score (standardised β 0.343, 95% confidence interval 0.188–0.498; FDR 4.21×10⁻⁴). Day-3 transcriptomic IFN-α and IFN-γ scores were associated with the corresponding Day-5 protein scores (FDR 0.0074 and 0.0235); no reverse-order model was FDR significant. Weighted and unweighted Day-5 Hallmark profiles were concordant (NES Spearman ρ=0.980), although positivity and residual covariate balance remained limited.",
            "Improved readability, standardised pathway names and retained every frozen numerical value. The patient/death notation is expanded once and then compressed.",
        ),
        (
            "Mortality-associated molecular rankings differed across landmark-specific survivor risk sets and between compartments. At the Day-5 landmark, the whole-blood transcriptomic ranking showed marked negative enrichment of cellular-maintenance programmes; later plasma proteomic rankings were dominated by tissue remodelling; and interferon pathways showed selective contemporaneous and forward cross-compartment associations. Independent validation is required, and these associations do not establish within-patient molecular trajectories, causal transfer or incremental prognostic value beyond clinical severity. Bulk whole-blood RNA-seq measurements could not distinguish cellular-compositional shifts from cell-intrinsic transcriptional change.",
            "Mortality-associated molecular rankings differed across landmark-specific survivor risk sets and between compartments. At the Day-5 landmark, the whole-blood transcriptomic ranking showed marked negative enrichment of cellular-maintenance programmes, whereas later plasma proteomic rankings were dominated by tissue remodelling. Interferon pathways showed selective contemporaneous and forward cross-compartment associations. Independent validation is required: these associations do not establish within-patient molecular trajectories, causal transfer or incremental prognostic value beyond clinical severity, and bulk whole-blood RNA-seq cannot distinguish changes in cellular composition from cell-intrinsic transcriptional change.",
            "Separated the main findings from the inferential limitations and removed a semicolon-heavy construction.",
        ),
        (
            "Previous analyses of the CMAISE resource have focused on baseline molecular heterogeneity, subgroup discovery or prediction.[10,11] Those studies established the value of the cohort but did not address the present question: whether mortality-associated pathway rankings differ across Day-1, Day-3 and Day-5 landmark-specific survivor risk sets in a fixed Day-1 SIC population, and whether selected pathways show reproducible associations across whole-blood and plasma compartments. Analysis of later samples also requires explicit attention to risk-set eligibility and sample observation, because patients who die before a sampling landmark cannot contribute a post-mortem molecular state and patients who survive may still differ in whether a sample is obtained.",
            "Previous analyses of the CMAISE resource focused on baseline molecular heterogeneity, subgroup discovery or prediction.[10,11] They did not determine whether mortality-associated pathway rankings differ across Day-1, Day-3 and Day-5 landmark-specific survivor risk sets in a fixed Day-1 SIC population, or whether selected pathways show reproducible associations across whole-blood and plasma compartments. Analyses of later samples also require explicit attention to risk-set eligibility and sample observation: patients who die before a sampling landmark cannot contribute a post-mortem molecular state, and survivors may differ systematically in whether a sample is obtained.",
            "Clarified the incremental question relative to previous CMAISE analyses and tightened the risk-set rationale.",
        ),
        (
            "We therefore analysed serial whole-blood RNA-seq and plasma protein measurements from a fixed Day-1 SIC cohort using separate landmark-specific survival models with delayed entry. Our objectives were to define molecular programmes associated with subsequent 60-day mortality at each sampling landmark, compare prognostic pathway architecture between compartments, identify reproducible contemporaneous and forward cross-omic associations, and evaluate the robustness of Day-5 RNA-seq findings to observed sampling selection. We hypothesised that mortality-associated molecular organisation would differ across landmarks and compartments, with only selected host-response programmes showing reproducible cross-omic association.",
            "We therefore analysed serial whole-blood RNA-seq and plasma protein measurements from a fixed Day-1 SIC cohort using separate landmark-specific survival models with delayed entry. We aimed to define molecular programmes associated with subsequent 60-day mortality at each landmark, compare prognostic pathway architecture between compartments, identify reproducible contemporaneous and forward cross-omic associations, and evaluate the robustness of Day-5 RNA-seq findings to observed sampling selection. We hypothesised that mortality-associated molecular architecture would differ across landmarks and compartments, with reproducible cross-omic associations confined to selected host-response programmes.",
            "Tightened the objectives and replaced the vague term ‘molecular organisation’ with ‘molecular architecture’.",
        ),
        (
            "Baseline variables were summarised as median [interquartile range] or number (percentage), with between-group separation described primarily using standardised mean differences. Arterial pH and calcium were excluded from Table 1 and from the exploratory clinical Cox analysis according to the prespecified rule for incomplete baseline observations. Before outcome review, the nine deposited infection-source levels were collapsed into six clinically defined groups: intestine was combined with abdomen; bloodstream and brain were combined with others/unknown; the lung/chest, urinary, soft-tissue and biliary/liver groups were retained. Exploratory univariable Cox models used the prespecified clinically interpretable increments reported in Supplementary Table S2. Proportional-hazards diagnostics and prespecified nonlinearity tests were performed, and the Benjamini–Hochberg procedure was applied across the displayed Cox contrasts. Infection source was evaluated using an overall likelihood-ratio test followed by category-specific contrasts with lung/chest infection as the reference. These analyses provided clinical context and were not used for covariate selection or construction of the molecular models.",
            "Baseline variables were summarised as median [interquartile range] or number (percentage), with between-group separation described primarily using standardised mean differences. Arterial pH and calcium were excluded from Table 1 and the exploratory clinical Cox analysis according to the prespecified rule for incomplete baseline observations. Before outcome review, the nine deposited infection-source levels were collapsed into six clinically defined groups: intestine was combined with abdomen; bloodstream and brain were combined with others/unknown; and lung/chest, urinary, soft-tissue and biliary/liver sources were retained (Supplementary Table S1). Exploratory univariable Cox models used the prespecified clinically interpretable increments reported in Supplementary Table S2. Proportional-hazards diagnostics and prespecified nonlinearity tests were performed, and the Benjamini–Hochberg procedure was applied across the displayed Cox contrasts. Infection source was evaluated using an overall likelihood-ratio test followed by category-specific contrasts with lung/chest infection as the reference. These analyses provided clinical context and were not used for covariate selection or construction of the molecular models.",
            "Added the new Supplementary Table S1 citation, which documents definitions, transformations, infection-source mapping and missing-variable exclusions without introducing new analyses.",
        ),
        (
            "Analyses were conducted in R using frozen scripts, package versions and random seeds recorded in the repository. The submission-freeze overlay reads formal result files and a single numerical-truth interface and does not refit primary models. Code, aggregate result tables, quality-assurance outputs and figure source data are available under the versioned repository tag analysis-freeze-v1.0-2026-07-13.[26] File-level SHA256 manifests were retained to verify that manuscript source data, figures and frozen result files were unchanged after closeout. Patient-level clinical and molecular data remain subject to CMAISE/OMIX controlled-access governance.[25]",
            "Analyses were conducted in R version 4.4.2 using frozen scripts, package versions and random seeds retained in the project archive. The submission-freeze overlay reads formal result files and a single numerical-truth interface and does not refit primary models. File-level SHA256 manifests verify that manuscript source data, figures and frozen outputs were unchanged during manuscript production. Participant-level clinical and molecular data remain subject to CMAISE/OMIX controlled-access governance.[25] Aggregate source data and reproducible code for the reported tables and figures are maintained in the project repository; the exact manuscript version will be archived in a citable release before submission.[26]",
            "Added the exact R version used for the frozen analysis and corrected the obsolete repository wording. A reviewer-accessible release and persistent archive identifier remain a mandatory author action before submission.",
        ),
        (
            "In exploratory univariable Cox analyses, 10 of 41 displayed contrasts retained BH-FDR <0.05. Higher SOFA score (HR 1.18 per point), lactate (HR 1.10 per mmol/L), sodium (HR 1.25 per 5 mmol/L), potassium (HR 1.42 per mmol/L), age (HR 1.22 per 10 years) and maximum heart rate (HR 1.11 per 10 beats/min), together with dementia (HR 3.20), were associated with higher 60-day mortality; a higher PaO2/FiO2 ratio (HR 0.83 per 50 mmHg) was associated with lower mortality. Infection source was associated with outcome overall (likelihood-ratio P=0.0038), with abdominal infection showing a lower hazard than lung/chest infection (HR 0.30, 95% CI 0.14–0.66; BH-FDR=0.0174). Blood urea nitrogen also retained BH-FDR significance but showed evidence of nonlinearity; its linear HR is therefore reported only as a summary contrast. These univariable estimates provide clinical context and were not interpreted as independent predictors or used to construct the molecular models (Supplementary Table S2).",
            "In exploratory univariable Cox analyses, 10 of 41 displayed contrasts retained BH-FDR <0.05. Higher SOFA score (HR 1.18 per point), lactate (HR 1.10 per mmol/L), sodium (HR 1.25 per 5 mmol/L), potassium (HR 1.42 per mmol/L), age (HR 1.22 per 10 years), maximum heart rate (HR 1.11 per 10 beats/min) and dementia (HR 3.20) were associated with higher 60-day mortality; a higher PaO2/FiO2 ratio (HR 0.83 per 50 mmHg) was associated with lower mortality. Infection source was associated with outcome overall (likelihood-ratio P=0.0038), with abdominal infection showing a lower hazard than lung/chest infection (HR 0.30, 95% CI 0.14–0.66; BH-FDR=0.0174). Blood urea nitrogen also retained BH-FDR significance but showed evidence of nonlinearity; its linear HR is therefore reported only as a summary contrast. These univariable estimates provide clinical context and were not interpreted as independent predictors or used to construct the molecular models (Supplementary Table S2 and Supplementary Figure A9).",
            "Added the supplementary forest-plot citation and clarified that the plot is a visual rendering of the same descriptive S2 contrasts, not an additional analysis.",
        ),
        (
            "Mortality-associated molecular rankings in patients with Day-1 identified SIC differed across Day-1, Day-3 and Day-5 landmark risk cohorts, and between whole blood and plasma compartments. Whole-blood rankings featured recurrent enrichment of heme metabolism, hypoxia and inflammatory pathways, alongside prominent negative enrichment of cellular maintenance programs at Day 5. By contrast, later plasma protein rankings were dominated by extracellular matrix and tissue remodelling signals. Interferon pathways exhibited reproducible contemporaneous and forward intercompartmental associations; however, these observational models did not establish causal intercompartmental signalling. Cell-type-resolved longitudinal validation incorporating synchronised profiling of SIC/DIC status, endothelial function, coagulation phenotypes and treatment exposures is required before inferring incremental prognostic value, specific cellular mechanisms or therapeutic relevance.",
            "Mortality-associated molecular rankings in patients with Day-1-defined SIC differed across Day-1, Day-3 and Day-5 landmark risk sets and between whole-blood and plasma compartments. Whole-blood rankings showed recurrent enrichment of heme metabolism, hypoxia and inflammatory pathways, together with prominent negative enrichment of cellular-maintenance programmes at Day 5. Later plasma protein rankings were dominated by extracellular-matrix and tissue-remodelling signals. Interferon pathways exhibited reproducible contemporaneous and forward cross-compartment associations, but these observational models did not establish causal intercompartmental signalling. Cell-type-resolved longitudinal validation with synchronised assessment of SIC/DIC status, endothelial function, coagulation phenotypes and treatment exposures is required before inferring incremental prognostic value, specific cellular mechanisms or therapeutic relevance.",
            "Standardised terminology and tightened the final conclusion without strengthening causal claims.",
        ),
        (
            "Individual-level clinical and molecular data are available through the controlled CMAISE/OMIX application process (OMIX011182) and are not redistributed in the public repository.[25] Analysis code, frozen aggregate result tables, numerical-truth interfaces, quality-assurance outputs and figure source data are available in the SIC-research repository under tag analysis-freeze-v1.0-2026-07-13.[26]",
            "Individual-level clinical and molecular data are controlled-access data available through the formal CMAISE/OMIX application process under accession OMIX011182; the authors are not authorised to redistribute them.[25] Aggregate result tables, figure source data, quality-assurance outputs and analysis code supporting the reported results are maintained at https://github.com/lvhao1123/SIC-longitudinal-multiomics. The exact manuscript version will be archived in a citable release before submission.[26]",
            "Corrected the governance statement and repository URL. The repository is currently private and has no verified manuscript release/tag; create a reviewer-accessible release and persistent archive identifier before submission, then replace the provisional final sentence.",
        ),
        (
            "Supplementary Table S2 presents the complete exploratory clinical univariable Cox analysis. Supplementary Tables S3 and S4 present complete RNA-seq gene-wise Cox/PH and Hallmark GSEA results; Supplementary Tables S5 and S6 present the corresponding plasma protein results; Supplementary Table S7 presents contemporaneous, forward and reverse-order cross-omic models; and Supplementary Table S8 presents the Day-5 RNA-seq availability/IPW diagnostics and Hallmark comparisons. Supplementary Figures A1–A8 present centre positivity, probability and weight distributions, covariate balance, Hallmark robustness, entry-boundary sensitivity and descriptive protein availability.",
            "Supplementary Table S1 documents clinical variable definitions, transformations, infection-source mapping and prespecified exclusions. Supplementary Table S2 presents the complete exploratory clinical univariable Cox analysis. Supplementary Table S3 presents the complete RNA-seq gene-wise Cox/PH results; Supplementary Table S4 presents the RNA-seq Hallmark GSEA and leading-edge results. Supplementary Table S5 presents the complete protein-wise Cox/PH results; Supplementary Table S6 presents the plasma protein Hallmark GSEA and leading-edge results. Supplementary Table S7 presents contemporaneous, forward and reverse-order cross-omic models, and Supplementary Table S8 presents the Day-5 RNA-seq availability/IPW diagnostics and Hallmark comparisons. Supplementary Figures A1–A8 present centre positivity, probability and weight distributions, covariate balance, Hallmark robustness, entry-boundary sensitivity and descriptive protein availability; Supplementary Figure A9 presents the exploratory clinical univariable Cox results as a forest plot.",
            "Completed the supplementary attachment map: S1–S8 and A1–A9 now have unique manuscript identifiers and independent upload files.",
        ),
        (
            "26. Lyu H, et al. SIC-research: frozen analysis code, aggregate results and source data. GitHub. https://github.com/lvhao1123/SIC-research/tree/analysis-freeze-v1.0-2026-07-13. Accessed 14 July 2026.",
            "26. Lyu H, et al. SIC longitudinal multi-omics: reproducible analysis code, aggregate results and figure source data. GitHub. https://github.com/lvhao1123/SIC-longitudinal-multiomics. Manuscript release and persistent archive identifier to be assigned before submission.",
            "Replaced the obsolete repository citation. This reference remains provisional until a public or reviewer-accessible frozen release is archived with a persistent identifier.",
        ),
        (
            "Supplementary Figures A1–A8 are presented below in numerical order. They provide the prespecified Day-5 RNA-seq availability/IPW diagnostics, Hallmark robustness analyses, entry-boundary sensitivity analysis and descriptive Day-5 plasma protein availability audit cited in the main text.",
            "Supplementary Figures A1–A9 are presented below in numerical order. Figures A1–A8 provide the prespecified Day-5 RNA-seq availability/IPW diagnostics, Hallmark robustness analyses, entry-boundary sensitivity analysis and descriptive Day-5 plasma protein availability audit. Figure A9 presents the complete displayed clinical univariable Cox contrasts as a supplementary forest plot.",
            "Extended the embedded supplementary-figure section to include A9.",
        ),
    ]

    for old, new, comment in replacements:
        replace_paragraph(doc, old, new, comment, annotated)

    standardise_table1_labels(doc, annotated)

    # Split the original very long pathway-score paragraph into three coherent
    # paragraphs without changing its content or statistical claims.
    long_panel = next(
        p for p in doc.paragraphs
        if p.text.startswith("For each sample and molecular layer, detected features were converted")
    )
    long_panel.clear()
    first = long_panel.add_run(
        "For each sample and molecular layer, detected features were converted to fractional within-sample ranks. "
        "For each pathway, the score was the mean fractional rank of its detected members minus 0.5; a score was "
        "calculated only when at least 10 pathway members were detected. Scores were standardised within each "
        "sampling landmark before cross-omic modelling. These rank-based scores were relative within-sample "
        "composite summaries rather than reflective latent scales. Rank transformation reduced dependence on "
        "absolute assay scale but did not remove centre, batch or cell-composition effects."
    )
    second = insert_paragraph_after(
        long_panel,
        "The analysis-prespecified core panel was frozen before patient-level contemporaneous and cross-lag "
        "modelling, but it was not prospectively registered before the single-omic outcome analyses. Pathways "
        "were not selected according to cross-omic model P values. The panel was intended to reduce multiplicity "
        "and preserve a clinically interpretable cross-omic hypothesis space, not to provide an exhaustive "
        "catalogue of SIC biology.",
    )
    third = insert_paragraph_after(
        second,
        "The 15 pathways represented four SIC-relevant domains supported by the reference study and consensus "
        "pathobiology: inflammatory and interferon host response (TNF-α signalling via NF-κB, IL-6/JAK/STAT3 "
        "signalling, inflammatory response, IFN-α response and IFN-γ response); thromboinflammation and "
        "erythroid/heme biology (coagulation, complement and heme metabolism); barrier and tissue remodelling "
        "(epithelial–mesenchymal transition, TGF-β signalling and apical junction); and hypoxic, oxidative and "
        "bioenergetic stress (hypoxia, oxidative phosphorylation, reactive oxygen species pathway and glycolysis). "
        "All 15 pathways met the prespecified coverage requirement; after quality control, each pathway contained "
        "at least 20 measured proteins and 48 transcriptomic features.",
    )
    if annotated:
        doc.add_comment(first, "Split a 250-word methods paragraph into three units: score construction, panel-freeze rationale and biological-domain/coverage rationale. No analysis or pathway membership was changed.", REVIEW_AUTHOR, REVIEW_INITIALS)
        comment_paragraph(doc, second, "The wording now makes the timing of panel specification explicit and avoids implying prospective registration.", annotated)
        comment_paragraph(doc, third, "The rationale for the 15 pathways is now stated as a prespecified, four-domain, coverage-qualified hypothesis space.", annotated)

    figure_titles = {
        1: "Study design, longitudinal risk sets and Day-5 availability estimand",
        2: "Time-specific whole-blood transcriptomic prognostic programmes",
        3: "Time-specific plasma protein prognostic programmes",
        4: "Pathway-selective contemporaneous and forward cross-omic associations",
    }
    for figure_number, short_title in figure_titles.items():
        bind_main_figure_block(
            doc,
            figure_number,
            short_title,
            main_figures[figure_number],
            width_inches=6.2,
        )
    legend_headings = [p for p in doc.paragraphs if p.text == "Figure legends"]
    if len(legend_headings) != 1:
        raise RuntimeError(f"Expected one Figure legends heading; found {len(legend_headings)}")
    remove_paragraph(legend_headings[0])

    # Replace the existing A7 rendering with the margin-corrected production
    # export. This is a figure-interface update only; source data and estimates
    # are unchanged.
    replace_embedded_figure(
        doc,
        "Supplementary Figure A7 |",
        a7,
        width_inches=6.25,
    )

    if not a9.exists():
        raise FileNotFoundError(a9)
    caption = doc.add_paragraph(style="Figure note")
    caption.paragraph_format.page_break_before = True
    caption.paragraph_format.keep_with_next = True
    caption_run = caption.add_run(
        "Supplementary Figure A9 | Exploratory clinical univariable Cox associations with 60-day mortality. "
        "Points and horizontal bars show hazard ratios and 95% confidence intervals from separate univariable "
        "Cox models in 504 patients with 84 deaths. Continuous variables are shown using the prespecified "
        "increments indicated in each label. aPTT, activated partial thromboplastin time; BMI, body mass index; "
        "BUN, blood urea nitrogen; COPD, chronic obstructive pulmonary disease; CRP, C-reactive protein; INR, "
        "international normalised ratio; MAP, mean arterial pressure; SBP, systolic blood pressure; SOFA, "
        "Sequential Organ Failure Assessment; WBC, white blood cell. The Benjamini–Hochberg procedure was applied "
        "across all 41 displayed "
        "contrasts; * indicates BH-FDR<0.05. † denotes nominal evidence against the proportional-hazards assumption "
        "and ‡ denotes nominal evidence of nonlinearity. Infection source was additionally associated with outcome "
        "in an overall likelihood-ratio test (P=0.0038). These univariable estimates are descriptive and are not "
        "mutually adjusted or causal."
    )
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(a9), width=Inches(6.5))
    if annotated:
        doc.add_comment(
            caption_run,
            text="A9 is intentionally supplementary. It improves transparent clinical reporting and readability but does not add independent-predictor evidence or strengthen the molecular claims.",
            author=REVIEW_AUTHOR,
            initials=REVIEW_INITIALS,
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--fig1", type=Path, required=True)
    parser.add_argument("--fig2", type=Path, required=True)
    parser.add_argument("--fig3", type=Path, required=True)
    parser.add_argument("--fig4", type=Path, required=True)
    parser.add_argument("--a7", type=Path, required=True)
    parser.add_argument("--a9", type=Path, required=True)
    parser.add_argument("--clean", type=Path, required=True)
    parser.add_argument("--annotated", type=Path, required=True)
    args = parser.parse_args()
    main_figures = {1: args.fig1, 2: args.fig2, 3: args.fig3, 4: args.fig4}
    apply_revision(args.source, args.clean, main_figures, args.a7, args.a9, annotated=False)
    apply_revision(args.source, args.annotated, main_figures, args.a7, args.a9, annotated=True)
    print(args.clean)
    print(args.annotated)


if __name__ == "__main__":
    main()
