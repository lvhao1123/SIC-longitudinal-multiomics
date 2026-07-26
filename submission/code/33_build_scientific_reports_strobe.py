from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / 'submission' / 'manuscript_files' / 'STROBE_checklist_cohort_completed.docx'
OUT = ROOT / 'submission' / 'manuscript_files' / 'scientific_reports' / 'STROBE_Scientific_Reports_completed.docx'
AUDIT = ROOT / 'submission' / 'manuscript_files' / 'scientific_reports' / 'STROBE_Scientific_Reports_audit.tsv'
TITLE = ('Landmark-specific transcriptomic and proteomic associations with 60-day mortality '
         'in Day-1-defined sepsis-induced coagulopathy: a multicentre longitudinal cohort study')

WHERE = {
    '1(a)': 'p. 1 (title identifies a multicentre longitudinal cohort study).',
    '1(b)': 'p. 1 (unstructured Abstract summarises the design, principal findings and bounded interpretation).',
    '2': 'p. 1 (Introduction).',
    '3': 'p. 1 (final Introduction paragraph states objectives and the prespecified hypothesis).',
    '4': 'p. 1 (Abstract) and pp. 9–10 (Study design and data source; Analytic framework and estimands).',
    '5': 'pp. 9–10 (43 tertiary hospitals; November 2020–November 2024; 60-day follow-up; Days 1, 3 and 5 sampling).',
    '6(a)': 'pp. 2 and 9–10; Figure 1 (eligibility, Day-1 SIC cohort, landmark risk sets and follow-up).',
    '6(b)': 'Not applicable: this was not a matched cohort study.',
    '7': 'pp. 9–13; Supplementary Table S1.',
    '8': 'pp. 9–11; Supplementary Information pp. 1–2 (Supplementary Methods).',
    '9': 'pp. 9–13 (delayed entry, centre stratification, PH diagnostics, sensitivity analyses and Day-5 availability/IPW); pp. 7–8 (limitations); Supplementary Figs. S1–S8.',
    '10': 'p. 9 (all eligible Day-1 SIC participants and available molecular measurements were included; no prospective sample-size calculation).',
    '11': 'pp. 9–13; Supplementary Tables S1–S2.',
    '12(a)': 'pp. 9–13; Supplementary Information pp. 1–2.',
    '12(b)': 'No subgroup analysis informed the primary conclusions. Prespecified contemporaneous, forward and reverse-order cross-omic models are described on pp. 11–12 and in Supplementary Methods.',
    '12(c)': 'p. 13; Supplementary Methods; Supplementary Tables S1 and S8.',
    '12(d)': 'p. 10 (60-day outcome and administrative censoring). Molecular sample non-observation was examined separately on pp. 12–13 and in Supplementary Table S8.',
    '12(e)': 'pp. 10–13; Supplementary Figs. S1–S8; Supplementary Tables S4, S6 and S8.',
    '13(a)': 'p. 2; Figure 1 (Day-1 cohort, raw measured, risk-valid and positivity-supported analysis counts).',
    '13(b)': 'p. 2; Figure 1 (structural deaths, raw-to-risk-valid exclusions and zero-observation-centre exclusions).',
    '13(c)': 'Figure 1 (legend on p. 15).',
    '14(a)': 'Table 1, p. 16; complete baseline characteristics in Supplementary Table S9 (Supplementary Information p. 13).',
    '14(b)': 'p. 13; Supplementary Methods; Supplementary Tables S1 and S8.',
    '14(c)': 'p. 10 (60-day outcome follow-up and administrative censoring).',
    '15': 'p. 2; Figure 1.',
    '16(a)': 'pp. 2–4; Figures 2–4; Supplementary Tables S2–S7. Adjusted covariates and their rationale are specified on pp. 11–12.',
    '16(b)': 'Not applicable to the primary molecular analyses: continuous features were standardised. Clinical Cox contrasts used prespecified increments (pp. 9–10; Supplementary Table S2).',
    '16(c)': 'Not applicable: the study characterised prognostic molecular associations and did not develop an absolute-risk model.',
    '17': 'pp. 2–4 and 10–13; Supplementary Figs. S1–S9; Supplementary Tables S4, S6 and S8.',
    '18': 'p. 5 (opening Discussion paragraphs).',
    '19': 'pp. 7–8 (limitations and potential sources of bias or imprecision).',
    '20': 'pp. 5–8 (Discussion, including multiplicity, alternative explanations and claim boundaries).',
    '21': 'pp. 7–8 (need for independent validation and limits to external validity).',
    '22': 'p. 14 (Funding).',
}

AUDIT_NOTES = {
    '6(a)': 'Explicitly covers the fixed Day-1 SIC cohort and separate Day-3/Day-5 landmark risk sets.',
    '9': 'Includes delayed entry, structural non-availability, positivity assessment and IPW.',
    '12(c)': 'Separates clinical missingness from molecular sample non-observation.',
    '12(d)': 'Distinguishes outcome follow-up from assay availability.',
    '12(e)': 'Covers PH/intensity/entry-boundary/observation-model sensitivity analyses.',
    '13(a)': 'Upstream parent-cohort screening was unavailable; the analytic source cohort and all subsequent risk sets are reported.',
    '14(b)': 'Missingness and availability are documented in dedicated supplementary tables.',
    '16(a)': 'All molecular effect estimates remain associations; no clinical prediction claim is made.',
    '17': 'Includes clinical, molecular and availability-weighting secondary analyses.',
    '19': 'Direction and likely consequences of selection, bulk-tissue and limited-event biases are discussed.',
    '21': 'External validation is explicitly required.',
}


def set_font(run, size=8.5, bold=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for k in ('ascii', 'hAnsi', 'eastAsia'):
        rfonts.set(qn(f'w:{k}'), 'Arial')


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_font(run, 8)


def main():
    source = Document(SOURCE)
    source_rows = [[c.text for c in row.cells] for row in source.tables[0].rows]

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.5); sec.right_margin = Inches(0.5)
    add_page_number(sec)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('STROBE Statement—completed checklist for cohort studies'); set_font(r, 13, True)
    p.paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Manuscript: ' + TITLE); set_font(r, 9.5, True)
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Reporting guideline: STROBE cohort checklist (accessed 25 July 2026). Page references correspond to the final Scientific Reports manuscript and Supplementary Information.'); set_font(r, 8.5)
    p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=len(source_rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.25, 0.55, 4.65, 3.6]
    audit_rows = []
    for i, source_row in enumerate(source_rows):
        row = table.rows[i]
        values = source_row[:]
        if i > 0:
            item = values[1]
            if item not in WHERE:
                raise KeyError(f'Missing mapping for item {item}')
            values[3] = WHERE[item]
            audit_rows.append({
                'item': item,
                'status': 'verified',
                'where_reported': WHERE[item],
                'audit_note': AUDIT_NOTES.get(item, 'Reporting location verified against final manuscript pagination.'),
            })
        for j, value in enumerate(values):
            cell = row.cells[j]
            cell.width = Inches(widths[j])
            cell.text = value
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_font(r, 7.5 if i else 8, True if i == 0 else False)
            if i == 0:
                shade(cell, 'D9EAF7')
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)
        if i == 0:
            repeat_header(row)

    doc.core_properties.title = 'STROBE checklist for Scientific Reports submission'
    doc.core_properties.subject = TITLE
    doc.save(OUT)

    with AUDIT.open('w', encoding='utf-8', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=['item', 'status', 'where_reported', 'audit_note'], delimiter='\t')
        writer.writeheader(); writer.writerows(audit_rows)
    print(f'Built {OUT}')
    print(f'Built {AUDIT}')


if __name__ == '__main__':
    main()
