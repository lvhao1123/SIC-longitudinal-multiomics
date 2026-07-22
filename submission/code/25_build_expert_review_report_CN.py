"""Create the three-reviewer Critical Care pre-submission report in Chinese."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Microsoft YaHei"
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.name = "Microsoft YaHei"

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SIC纵向多组学论文：Critical Care投稿前专家评估")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("基于v2.0文稿、冻结统计结果、4张主图及S1–S8/A1–A9投稿附件").italic = True

    doc.add_heading("输入范围、边界与总体判断", level=1)
    doc.add_paragraph(
        "本次评估以现有冻结结果为边界，不重新拟合Cox、IPW、GSEA或跨组学模型。评估对象包括："
        "504例Day-1定义SIC队列的临床描述；Day-1、Day-3和Day-5风险集有效的全血RNA与血浆蛋白"
        "预后关联；Hallmark GSEA；Day-5 availability/IPW敏感性分析；同期、前向及反向跨组学模型；"
        "主图1–4；Supplementary Tables S1–S8与Supplementary Figures A1–A9。"
    )
    p = doc.add_paragraph()
    p.add_run("总体判断：").bold = True
    p.add_run(
        "文章已具备向Critical Care投稿的统计与叙事基础，研究问题与期刊关注的重症临床转化、"
        "机制分层和多中心证据高度匹配。当前合理定位是‘有竞争力但仍需作者完成投稿接口动作’，"
        "而不是继续增加探索性模型。主要剩余风险为独立验证缺失、晚期蛋白事件数有限、D3/D5缺乏同期"
        "凝血表型、以及代码仓库尚未形成可供审稿访问的永久冻结版本。Critical Care官方页面强调"
        "高质量、循证且能促进危重患者照护的研究，本稿的临床意义必须始终通过谨慎的关联性表述来体现："
        "https://link.springer.com/journal/13054。"
    )

    doc.add_heading("Reviewer 1：重症医学与临床转化", level=1)
    doc.add_paragraph(
        "总体评价：研究聚焦SIC这一临床可识别且具有动态病理生理的重症综合征，纵向多组学设计优于"
        "单一Day-1横断面研究。文章最有价值的临床信息不是某个单基因或单蛋白，而是不同采样landmark"
        "后续死亡风险关联的分子结构发生重排，以及全血与血浆区室提供互补信息。"
    )
    doc.add_heading("主要意见", level=2)
    add_numbered(doc, [
        "必须持续强调Day-3和Day-5分析是相应landmark存活者中的独立时间特异性预后模型，不能将其叙述为同一模型内的time-dependent covariate，也不能将不同风险集的NES直接解释为患者内轨迹。v2.1已在摘要、Methods、Results和图注中统一这一边界。",
        "文章无法直接证明SIC由高凝向低凝转化，也无法判定Day-5患者仍满足SIC或已进展为显性DIC。原因是缺乏D3/D5同期SIC/DIC评分、血小板、PT/INR、纤维蛋白原、D-二聚体、黏弹力试验及治疗暴露。Discussion中的临床外推应保持‘与动态血栓炎症框架相容’，不能写成‘证实凝血表型演变’。",
        "Day-5蛋白组仅14个后续死亡事件，ECM/tissue-remodelling结果可作为通路层面的假设生成证据，但不应升格为可用于个体预测或治疗选择的生物标志物。当前文稿已明确这一限制。",
        "临床单因素Cox森林图有展示价值，但不应进入正文主图。其作用是提高临床可读性、公开全部41个预设对比并显示PH/非线性警示；它不会把单因素变量变成独立预测因子，也不会给分子机制结论增加统计证据。将其作为Supplementary Figure A9是合适的。",
    ])
    doc.add_heading("次要意见与已实施修改", level=2)
    add_bullets(doc, [
        "Table 1保留504例总体、420例存活/删失和84例死亡分层；动脉pH和血钙因预设缺失规则未纳入，符合作者决定。",
        "感染来源六分类应明确为结局查看前完成的临床整合；Supplementary Table S1现已保存映射及处理依据。",
        "标题已缩短，摘要和结论已减少‘机制驱动’色彩，突出landmark-specific mortality association。",
        "建议投稿信突出风险集处理、中心分层、PH敏感性、蛋白全局强度审计和availability/IPW边界，而非宣传大量生物标志物。",
    ])
    doc.add_paragraph("Reviewer 1建议：Major revision后可送审；所需修订主要为叙事和外部验证边界，而非新增统计分析。")

    doc.add_heading("Reviewer 2：生存统计、多组学与结果稳健性", level=1)
    doc.add_paragraph(
        "总体评价：统计框架整体成熟。延迟进入、中心分层、完整Cox Wald排序、完整Hallmark家族FDR、"
        "PH-pass敏感性、蛋白全局强度敏感性、正反向跨组学模型和Day-5观察概率审计共同构成了较强的"
        "方法学防线。最重要的是，作者没有用显著基因先筛选再做GSEA。"
    )
    doc.add_heading("主要意见", level=2)
    add_numbered(doc, [
        "分子Cox模型为centre-stratified但未调整临床严重程度，因此估计的是总预后关联，而不是severity-independent或增量预后价值。当前Methods和Limitations已正确写明；摘要与讨论中不得使用‘independent prognostic factor’或‘improved prediction’。",
        "三个时间点样本量和事件数不同，NES的跨时间比较属于不同风险集中死亡相关排序的比较，而非同一统计量的纵向斜率。图2和图3的星号仅代表显示的主要分析BH-FDR，PH-pass和global-intensity结果必须继续放在补充表与source data中，不能暗示星号编码了联合稳健性。当前图注已符合这一规则。",
        "Day-5 availability/IPW结果必须保持敏感性分析定位。最低概率0.00505、最大原始权重20.85、ESS/observed N=0.913但最大加权绝对SMD=0.450，支持‘总体通路高度一致但positivity和残余平衡受限’，不能写成‘IPW消除了选择偏倚’。",
        "跨组学forward模型在控制前一时间点蛋白评分后显示IFN-α/γ关联，反向模型未达FDR显著；这一时间不对称不能证明RNA导致蛋白变化、翻译滞后或中介作用。现有表述‘cross-compartment temporal association’恰当。",
        "15条核心Hallmark通路不是前瞻注册，但在患者级跨组学建模前冻结，且没有按跨组学P值筛选。Methods现已拆分长段落，明确四个生物学域、覆盖阈值及减少多重性的目的，足以回应选择依据质疑。",
    ])
    doc.add_heading("风险与建议", level=2)
    add_bullets(doc, [
        "不建议追加WGCNA、无监督聚类、机器学习或更多数据库富集；这些分析会扩张问题而非解决当前推断限制。",
        "建议在补充材料中保留全部gene/protein Cox与PH结果、完整Hallmark和leading edge、全部六个IPW情景及反向跨组学模型；S3–S8已独立封装。",
        "临床单因素森林图应显示所有41个对比、BH-FDR及诊断标记，并在图注中明确未经互相调整；A9已按该规范生成。",
    ])
    doc.add_paragraph("Reviewer 2建议：统计分析充分，可冻结；新增探索性模型的边际收益低于其多重性和过拟合代价。")

    doc.add_heading("Reviewer 3：可复现性、图表与出版规范", level=1)
    doc.add_paragraph(
        "总体评价：投稿生产层已接近完整。4张主图围绕队列/风险集、RNA、Protein和跨组学依次展开，"
        "比大量火山图或逐通路富集曲线更适合Critical Care读者。S1–S8与A1–A9提供了可审计的机器可读"
        "附件，且不含PatientID或个体权重。"
    )
    doc.add_heading("主要意见", level=2)
    add_numbered(doc, [
        "受控数据不需要公开再分发。正确的数据声明是：个体层面临床和组学数据须经NGDC/OMIX011182正式申请获得，作者无权在GitHub再发布；公开仓库仅提供代码、聚合结果、图形source data和QA。v2.1已修正。",
        "新仓库https://github.com/lvhao1123/SIC-longitudinal-multiomics目前为私有开发仓库，且本次审计未发现可核验的正式manuscript release/tag。因此旧SIC-research链接和analysis-freeze-v1.0-2026-07-13声明已删除。投稿前必须创建审稿可访问的冻结release，并优先通过Zenodo等获得永久DOI；这是当前唯一明确的可复现性阻断项。",
        "S1–S8应作为8个独立xlsx附件上传，A1–A9应统一置于Supplementary data或分别上传，不能把补充图穿插在正文结果段落。文稿末尾保留补充图是可接受的投稿草稿形式；实际投稿界面按期刊要求拆分。",
        "图2和图3的正/负NES临床意义及主要分析FDR星号已经统一。图4的六panel需要整页或双栏宽度，当前分组标题、rho/beta解释和CI可辨识；不要为了单栏适配删减既有冻结结果。",
        "AI使用声明已限定为代码审阅、实施支持、组织和语言辅助，并由作者承担验证责任；投稿前应再次核对Springer Nature当时生效的生成式AI披露要求。",
    ])
    doc.add_heading("语言与图表评价", level=2)
    add_bullets(doc, [
        "v2.1已统一programme、centre、normalised等英式拼写，以及TNF-α signalling via NF-κB、IL-6/JAK/STAT3 signalling、MYC targets V1、E2F targets、DNA repair和mTORC1 signalling等通路名称。",
        "批注版包含19条真实Word评论，标出标题、摘要、长Methods段落、S1/A9接口及仓库声明的修改理由；clean版不含评论。",
        "A9的横轴、HR/CI、FDR和PH/非线性标记均完整显示；它应作为补充透明度图，而非‘为文章加分’的装饰性图。",
    ])
    doc.add_paragraph("Reviewer 3建议：完成永久release/DOI、最终参考文献核验和投稿附件命名后即可提交。")

    doc.add_heading("跨审稿人综合意见", level=1)
    doc.add_heading("共同核心主张", level=2)
    doc.add_paragraph(
        "三位审稿人均认为最可辩护的主张是：在Day-1定义SIC队列中，后续死亡相关的全血转录和血浆蛋白"
        "通路结构具有landmark和区室特异性；晚期全血排序更突出细胞维持程序的负富集，血浆蛋白更突出"
        "ECM/组织重塑，而IFN通路表现出选择性的同期和前向跨区室关联。"
    )
    doc.add_heading("共同认可的证据", level=2)
    add_bullets(doc, [
        "504例Day-1 SIC固定队列及明确的D1/D3/D5风险集；",
        "中心分层Cox、PH审计、完整Cox排序和全Hallmark家族BH-FDR；",
        "蛋白QC与全局强度敏感性；",
        "Day-5 availability/IPW的positivity、权重、平衡和边界审计；",
        "同期、forward和reverse跨组学结果的统一展示；",
        "聚合source data与代码层的可追溯性。",
    ])
    doc.add_heading("共同认为仍缺失的证据", level=2)
    add_bullets(doc, [
        "独立外部纵向多组学验证；",
        "D3/D5同期SIC/DIC和凝血功能表型；",
        "精确采样时间、未采样原因和时间变化治疗暴露；",
        "细胞组成/单细胞验证及直接内皮、NET、凝血酶生成、纤溶、游离heme和线粒体功能标志物；",
        "可供审稿访问且带永久标识符的代码release。",
    ])

    doc.add_heading("最终发表可行性与作者动作", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["项目", "当前状态", "是否阻断写作", "建议"]):
        set_cell(cell, text, True)
    rows = [
        ("正式统计分析", "已充分并冻结", "否", "不新增探索性模型"),
        ("正文语言与推断边界", "v2.1已修补", "否", "以clean版继续作者核对"),
        ("Table 1", "结构与既定排除规则一致", "否", "保留pH/钙排除说明"),
        ("临床单因素森林图", "A9已生成", "否", "仅作补充图"),
        ("S1–S8/A1–A9附件", "已独立编号和封装", "否", "投稿时按manifest上传"),
        ("数据治理", "受控OMIX申请", "否", "不公开个体数据"),
        ("代码release/DOI", "尚未最终归档", "是：投稿前", "建立冻结release和永久标识符"),
        ("外部验证", "缺失", "不阻断本稿，但降低优先级", "在局限中明确并作为后续研究"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell(cell, text)

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.add_run("最终建议：").bold = True
    final.add_run(
        "可以正式进入Critical Care投稿准备，不需要继续增加主体数据分析。完成代码冻结release/DOI、"
        "参考文献最终核验、作者/基金/伦理信息确认和投稿系统附件拆分后即可提交。若送外审，最可能的决定"
        "仍是major revision，核心质疑将集中在外部验证、后期临床表型缺失和事件数，而不是现有统计流程不足。"
    )

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(args.out)


if __name__ == "__main__":
    main()
