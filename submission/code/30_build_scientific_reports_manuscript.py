from __future__ import annotations

import csv
import re
import shutil
import zipfile
from collections import OrderedDict
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / 'submission' / 'manuscript_files'
WORK = ROOT / 'submission' / 'manuscript_support'
OUT = BASE / 'scientific_reports'
OUT.mkdir(parents=True, exist_ok=True)

SOURCE_MANUSCRIPT = BASE / 'JIC_manuscript_clean.docx'
SOURCE_SUPP = BASE / 'Additional_file_1_Supplementary_methods_and_figures.docx'
SOURCE_STROBE = BASE / 'STROBE_checklist_cohort_completed.docx'

TITLE = ('Landmark-specific transcriptomic and proteomic associations with 60-day mortality '
         'in Day-1-defined sepsis-induced coagulopathy: a multicentre longitudinal cohort study')
AUTHORS = 'Hao Lyu¹, Lingwei Li², Jianxing Guo¹, Dong Zhang¹*'
AFFILIATIONS = [
    '¹ Department of Critical Care Medicine, The First Hospital of Jilin University, No. 1 Xinmin Street, Changchun, Jilin 130021, China.',
    '² Department of Cadre Ward, The First Hospital of Jilin University, No. 1 Xinmin Street, Changchun, Jilin 130021, China.',
]
CORRESPONDING = '*Correspondence: Dong Zhang (zhangdong@jlu.edu.cn)'
EMAILS = ('Hao Lyu: haolv25@mails.jlu.edu.cn; Lingwei Li: lilw24@mails.jlu.edu.cn; '
          'Jianxing Guo: guojianxing@jlu.edu.cn')
RELEASE_URL = ('https://github.com/lvhao1123/SIC-longitudinal-multiomics/'
               'releases/tag/scientific-reports-submission-v1.1')

ABSTRACT = (
    'Sepsis-induced coagulopathy (SIC) reflects dysregulated interactions among inflammation, '
    'endothelial injury and haemostasis, but the molecular features associated with mortality may '
    'differ across illness stages. We analysed 504 adults with Day-1 SIC from a multicentre Chinese '
    'cohort using separate Day-1, Day-3 and Day-5 landmark Cox models for 60-day mortality applied '
    'to whole-blood RNA sequencing and plasma proteomics. Signed feature rankings underwent Hallmark '
    'enrichment, selected pathway scores were examined across omic layers, and Day-5 RNA-seq results '
    'were tested with positivity-supported inverse-probability weighting. Risk-valid RNA-seq sets '
    'comprised 504/84, 420/67 and 320/53 patients/deaths; proteomic sets comprised 168/27, 147/18 '
    'and 114/14. Heme, hypoxia and inflammatory programmes were repeatedly enriched in '
    'mortality-associated transcriptomic rankings, whereas Day-5 rankings showed strong negative '
    'enrichment of oxidative phosphorylation, MYC, unfolded-protein-response, E2F and DNA-repair '
    'programmes. Later proteomic rankings were dominated by extracellular-matrix and tissue-remodelling '
    'signals. Interferon pathways showed selective contemporaneous and forward cross-omic associations. '
    'Weighted and unweighted Day-5 Hallmark profiles were highly concordant (Spearman ρ=0.980), although '
    'positivity and residual balance were limited. These landmark-specific associations define stage- '
    'and compartment-dependent prognostic biology but require independent validation and do not establish '
    'causal mechanisms or within-patient molecular trajectories.'
)
KEYWORDS = ('Sepsis-induced coagulopathy; Multi-omics; Landmark analysis; Transcriptomics; '
            'Proteomics; Mortality')

AI_DISCLOSURE = (
    'ChatGPT (OpenAI) was used as an assistive tool during manuscript drafting, restructuring, '
    'language editing, and the development and review of analysis scripts. All analyses were executed '
    'through the version-controlled workflow, and the reported numerical results, tables, and figures '
    'were verified against the frozen outputs and predefined quality-assurance checks. The authors '
    'retained final responsibility for the study design, analytical choices, interpretation, source '
    'verification, and conclusions; critically reviewed and revised all AI-assisted material; and '
    'approved the final manuscript. Generative AI was not used to create or modify the scientific figures.'
)

DATA_AVAILABILITY = (
    'Individual-level clinical, transcriptomic and proteomic data are controlled-access data available '
    'through the formal CMAISE/OMIX application process under accession OMIX011182. The authors are not '
    'authorised to redistribute these participant-level data. Aggregate non-identifiable result tables, '
    'figure source data and quality-assurance materials supporting this study are available in the '
    f'versioned repository release at {RELEASE_URL}.'
)

CODE_AVAILABILITY = (
    'Version-controlled analysis code, software-environment records, numerical-truth interfaces and '
    'reproducibility tests for the reported analyses are archived in the immutable repository release '
    f'at {RELEASE_URL}. The repository does not contain participant-level clinical or molecular data.'
)

ETHICS = (
    'This study was conducted in accordance with the Declaration of Helsinki and applicable regulations. '
    'The parent study was approved by the Ethics Committee of Sir Run Run Shaw Hospital '
    '(Approval No. 20201014-39), with additional local ethics approvals obtained at participating centres '
    'as required. Written informed consent was obtained from all participants or their legally authorised representatives.'
)

TRANSLATIONAL = (
    'Clinical translation should proceed through staged validation rather than immediate construction of '
    'a static prognostic panel. Prospective multicentre studies should prespecify sampling windows, '
    'reconstruct contemporaneous SIC and overt disseminated intravascular coagulation status, record '
    'treatment exposures, and test whether a reduced RNA/protein panel improves prognostic performance '
    'beyond contemporaneous severity and coagulation measures. Cell-resolved profiling and direct '
    'measurements of endothelial injury, coagulation and cellular metabolism could test biological '
    'plausibility. Independently specified biomarker-by-treatment analyses or enrichment trials would '
    'still be required before any therapeutic inference. The present findings therefore do not support '
    'treatment selection.'
)


def load_paragraphs(path: Path) -> dict[int, str]:
    doc = Document(path)
    return {i: p.text for i, p in enumerate(doc.paragraphs)}


SRC = load_paragraphs(SOURCE_MANUSCRIPT)
SUPP_SRC = load_paragraphs(SOURCE_SUPP)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_header_and_nosplit(table):
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)


def set_run_font(run, size=10.5, bold=None, italic=None, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), 'Arial')


def style_paragraph(paragraph, size=10.5, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=0, space_after=6, line_spacing=1.15, keep_with_next=False):
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold if bold else None, italic=italic if italic else None)


def add_page_number(footer_paragraph):
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9)


def enable_line_numbering(section):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn('w:lnNumType'))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement('w:lnNumType')
    ln.set(qn('w:countBy'), '1')
    ln.set(qn('w:start'), '1')
    ln.set(qn('w:restart'), 'continuous')
    sect_pr.append(ln)


def configure_section(section, landscape=False, margins=(0.9, 0.9, 0.8, 0.8), line_numbers=True):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    top, bottom, left, right = margins
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    if line_numbers:
        enable_line_numbering(section)
    footer = section.footer
    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        p.clear()
    add_page_number(p)


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    for name, size in [('Title', 14), ('Heading 1', 12), ('Heading 2', 10.5), ('Heading 3', 10.5)]:
        s = styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
        s._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
        s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


def add_text_paragraph(doc, text: str, *, style=None, highlight=False, bold=False, italic=False,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, space_before=0, space_after=6,
                       keep_with_next=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold if bold else None, italic=italic if italic else None)
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    style_paragraph(p, size=size, bold=bold, italic=italic, alignment=alignment,
                    space_before=space_before, space_after=space_after,
                    keep_with_next=keep_with_next)
    return p


def add_heading(doc, text: str, level=1, highlight=False):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_run_font(r, size=12 if level == 1 else 10.5, bold=True)
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    style_paragraph(p, size=12 if level == 1 else 10.5, bold=True,
                    space_before=10 if level == 1 else 6, space_after=4,
                    keep_with_next=True)
    return p


def parse_citation_group(group: str) -> list[int]:
    nums: list[int] = []
    for part in group.split(','):
        part = part.strip()
        if not part:
            continue
        if '-' in part or '–' in part:
            a, b = re.split(r'[-–]', part)
            nums.extend(range(int(a), int(b) + 1))
        else:
            nums.append(int(part))
    return nums


def compress_numbers(nums: list[int]) -> str:
    nums = sorted(dict.fromkeys(nums))
    out = []
    i = 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 2:
            out.append(f'{nums[i]}–{nums[j]}')
        elif j - i == 1:
            out.extend([str(nums[i]), str(nums[j])])
        else:
            out.append(str(nums[i]))
        i = j + 1
    return ', '.join(out)


CIT_RE = re.compile(r'\[([0-9,\-– ]+)\]')


def get_citation_order(texts: Iterable[str]) -> list[int]:
    order: list[int] = []
    seen = set()
    for text in texts:
        for m in CIT_RE.finditer(text):
            for n in parse_citation_group(m.group(1)):
                if n not in seen:
                    seen.add(n)
                    order.append(n)
    return order


def remap_citations(text: str, mapping: dict[int, int]) -> str:
    def repl(m):
        old = parse_citation_group(m.group(1))
        new = [mapping[n] for n in old]
        return '[' + compress_numbers(new) + ']'
    return CIT_RE.sub(repl, text)


def parse_reference(raw: str):
    raw = re.sub(r'^\d+\.\s*', '', raw).strip()
    m = re.match(r'(?P<authors>.+?)\. (?P<title>.+?)\. (?P<journal>.+?) (?P<year>\d{4}), (?P<vol>[^:]+):(?P<pages>.+?)\. doi: (?P<doi>.+?)\.$', raw)
    if not m:
        raise ValueError(f'Could not parse reference: {raw}')
    return m.groupdict()


def initials_with_periods(initials: str) -> str:
    out = []
    for token in re.split(r'(-)', initials):
        if token == '-':
            out.append('-')
        else:
            out.append(' '.join(ch + '.' for ch in token if ch.isalpha()))
    return ''.join(out)


def format_authors(raw: str) -> str:
    raw = raw.replace(', et al', ', et al.')
    if 'et al' in raw:
        first = raw.split(',')[0].strip()
        # Last token is initials; the rest is surname.
        parts = first.split()
        surname = ' '.join(parts[:-1])
        initials = initials_with_periods(parts[-1])
        return f'{surname}, {initials} et al.'
    author_chunks = [a.strip() for a in raw.split(',')]
    formatted = []
    for chunk in author_chunks:
        parts = chunk.split()
        surname = ' '.join(parts[:-1])
        initials = initials_with_periods(parts[-1])
        formatted.append(f'{surname}, {initials}')
    if len(formatted) == 1:
        return formatted[0] + '.'
    if len(formatted) == 2:
        return formatted[0] + ' & ' + formatted[1] + '.'
    return ', '.join(formatted[:-1]) + ' & ' + formatted[-1] + '.'


JOURNAL_MAP = {
    'JAMA': 'JAMA',
    'BMJ Open': 'BMJ Open',
    'J Thromb Haemost': 'J. Thromb. Haemost.',
    'Inflamm Res': 'Inflamm. Res.',
    'Nat Med': 'Nat. Med.',
    'Sci Transl Med': 'Sci. Transl. Med.',
    'EBioMedicine': 'eBioMedicine',
    'Nat Commun': 'Nat. Commun.',
    'Thromb Res': 'Thromb. Res.',
    'Med Res': 'Med. Res.',
    'Scand J Stat': 'Scand. J. Stat.',
    'Proc Natl Acad Sci': 'Proc. Natl Acad. Sci. USA',
    'Cell Syst': 'Cell Syst.',
    'Stat Methods Med Res': 'Stat. Methods Med. Res.',
    'PLoS One': 'PLoS ONE',
    'Nat Rev Immunol': 'Nat. Rev. Immunol.',
    'Arterioscler Thromb Vasc Biol': 'Arterioscler. Thromb. Vasc. Biol.',
    'Nat Immunol': 'Nat. Immunol.',
    'Lancet Respir Med': 'Lancet Respir. Med.',
    'Ann Intensive Care': 'Ann. Intensive Care',
    'Front Immunol': 'Front. Immunol.',
    'Am J Respir Crit Care Med': 'Am. J. Respir. Crit. Care Med.',
}


def add_reference_paragraph(doc, number: int, ref: dict):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(f'{number}. {format_authors(ref["authors"])} ')
    set_run_font(r, size=9)
    r = p.add_run(ref['title'].rstrip('.') + '. ')
    set_run_font(r, size=9)
    journal = JOURNAL_MAP.get(ref['journal'], ref['journal'])
    r = p.add_run(journal + ' ')
    set_run_font(r, size=9, italic=True)
    volume = re.sub(r'\(.*?\)', '', ref['vol'])
    r = p.add_run(volume)
    set_run_font(r, size=9, bold=True)
    pages = ref['pages'].replace('-', '–')
    r = p.add_run(f', {pages} ({ref["year"]}). ')
    set_run_font(r, size=9)
    r = p.add_run(f'https://doi.org/{ref["doi"]}')
    set_run_font(r, size=9)
    return p


# Build section content from frozen manuscript paragraphs.
intro_ids = [29, 30, 31, 32]
result_blocks = [
    ('Cohort construction and longitudinal molecular availability', [72, 73, 74]),
    ('Whole-blood mortality-associated rankings differed across landmark-specific risk sets', [81, 82, 83]),
    ('Plasma protein mortality-associated programmes concentrated on tissue remodelling', [89]),
    ('Cross-omic concordance was pathway selective rather than global', [93]),
    ('Interferon pathways showed selective forward cross-compartment associations', [95, 96, 97]),
    ('Day-5 transcriptomic pathway patterns were largely preserved after availability weighting', [102, 103]),
]
discussion_ids = [107, 108] + list(range(110, 115)) + list(range(116, 119)) + list(range(120, 123)) + list(range(124, 127)) + [132, 133, 135]
method_blocks = [
    ('Study design and data source', [35]),
    ('Analytic framework and estimands', [37]),
    ('Participants, SIC definition and outcome', [39, 40]),
    ('Clinical descriptive and exploratory survival analyses', [42]),
    ('Longitudinal sample linkage and time-specific risk sets', [44]),
    ('Molecular data generation in the parent cohort', [46]),
    ('RNA-seq processing', [48]),
    ('Plasma protein processing and quality control', [50]),
    ('Time-specific molecular survival models', [52, 53]),
    ('Hallmark gene-set enrichment', [55]),
    ('Patient-level pathway scores and cross-omic analysis', [57, 58, 59, 60]),
    ('Day-5 RNA-seq availability and inverse-observation-probability weighting', [62, 63, 64]),
    ('Missing data, multiplicity and interpretation hierarchy', [66]),
]

# Approved edits. Scientific values are unchanged.
TEXT_REPLACEMENTS = {
    42: SRC[42] + ' A concise set of clinically central variables is shown in Table 1; complete baseline characteristics are provided in Supplementary Table S9.',
    72: SRC[72] + ' Selected baseline characteristics are shown in Table 1, with the complete baseline table provided as Supplementary Table S9.',
    73: SRC[73].replace('Table 1', 'Table 1 and Supplementary Table S9'),
    74: SRC[74],
}


def source_text(i: int) -> str:
    text = TEXT_REPLACEMENTS.get(i, SRC[i])
    text = re.sub(r'Supplementary Figure A([1-9])', r'Supplementary Fig. S\1', text)
    text = re.sub(r'Supplementary Figures A([1-9])–A([1-9])', r'Supplementary Figs. S\1–S\2', text)
    return text


# Build all body texts once to establish first-citation order after reordering.
pre_map_texts = []
pre_map_texts.extend(source_text(i) for i in intro_ids)
for _, ids in result_blocks:
    pre_map_texts.extend(source_text(i) for i in ids)
pre_map_texts.extend(source_text(i) for i in discussion_ids[:-1])
pre_map_texts.append(TRANSLATIONAL)
pre_map_texts.append(source_text(132))
pre_map_texts.append(source_text(133))
pre_map_texts.append(source_text(135))
for _, ids in method_blocks:
    pre_map_texts.extend(source_text(i) for i in ids)
pre_map_texts.append(ETHICS)
pre_map_texts.append(CODE_AVAILABILITY)
pre_map_texts.append(AI_DISCLOSURE)

citation_order = get_citation_order(pre_map_texts)
# Retain any reference not cited in first occurrence order at end for audit; normally none.
reference_paragraphs = [text for text in SRC.values() if re.match(r'^\d+\.\s+', text.strip())]
raw_refs = {}
for text in reference_paragraphs:
    match = re.match(r'^(\d+)\.\s+', text.strip())
    if match is None:
        continue
    raw_refs[int(match.group(1))] = text
if sorted(raw_refs) != list(range(1, 30)):
    raise RuntimeError(f'Expected references 1-29; found {sorted(raw_refs)}')
for n in sorted(raw_refs):
    if n not in citation_order:
        citation_order.append(n)
CITATION_MAP = {old: new for new, old in enumerate(citation_order, start=1)}


def mapped(text: str) -> str:
    return remap_citations(text, CITATION_MAP)


CHANGE_LOG = [
    ('Abstract', 'journal compliance', 'Converted the structured abstract to a single paragraph of no more than 200 words while retaining only frozen principal results.'),
    ('Keywords', 'journal compliance', 'Reduced the keyword list from eight to six indexing terms.'),
    ('Main-text structure', 'structural relocation', 'Reordered the manuscript to Introduction, Results, Discussion and Methods; moved the standalone conclusion into the final Discussion paragraph.'),
    ('Discussion', 'interpretive expansion', 'Clarified landmark-specific rather than within-patient interpretation and organised biological meaning by immune-inflammatory, metabolic-stress and tissue-remodelling domains.'),
    ('Discussion', 'claim restriction', 'Compressed translational implications and explicitly stated that the findings do not support treatment selection.'),
    ('Table 1', 'journal compliance', 'Retained a concise one-page table of clinically central variables; moved the complete original table to Supplementary Table S9 without changing values.'),
    ('Supplementary material', 'cross-reference renumbering', 'Renumbered Supplementary Figures A1–A9 as Supplementary Figs. S1–S9 in all submission-facing files.'),
    ('Methods', 'clarification', 'Inserted the complete ethics and informed-consent statement into Methods.'),
    ('Methods', 'journal compliance', 'Added a dedicated Code availability subsection and an explicit LLM-use disclosure consistent with current Scientific Reports policy.'),
    ('Data availability', 'journal compliance', 'Separated controlled participant-level data from publicly released aggregate outputs and reproducibility materials.'),
    ('References', 'journal compliance', 'Renumbered citations after section relocation and converted the bibliography to standard Nature style while preserving citation mapping.'),
]


def add_title_page(doc, highlighted=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, size=14, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = add_text_paragraph(doc, AUTHORS, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, space_after=8)
    for aff in AFFILIATIONS:
        add_text_paragraph(doc, aff, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=3)
    add_text_paragraph(doc, CORRESPONDING, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=3)
    add_text_paragraph(doc, EMAILS, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, space_after=12)


def add_main_table(doc: Document, highlighted=False):
    source_doc = Document(SOURCE_MANUSCRIPT)
    src_table = source_doc.tables[0]
    rows = [[cell.text for cell in row.cells] for row in src_table.rows]
    labels = [r[0] for r in rows]
    selected_labels = [
        'Age', 'Lactate', 'BUN', 'Platelet', 'INR', 'D-dimer', 'SOFA score', 'PF',
        'Sex, male, n (%)', 'Diabetes mellitus', 'Hypertension', 'Heart failure',
        'Cerebrovascular disease', 'Renal failure', 'Infection source, n (%)',
        'Abdomen', 'Biliary/liver', 'Lung/chest', 'Others/unknown', 'Soft tissue', 'Urinary'
    ]
    selected = [rows[0]] + [rows[labels.index(label)] for label in selected_labels]
    selected = [([row[0], '', '', '', '', ''] if row[0] == 'Infection source, n (%)' else row) for row in selected]

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, landscape=True, margins=(0.45, 0.45, 0.45, 0.45), line_numbers=True)
    add_text_paragraph(doc, 'Table 1. Selected baseline characteristics of the Day-1 SIC cohort according to 60-day outcome',
                       size=10.5, bold=True, space_after=4, keep_with_next=True,
                       highlight=highlighted)
    table = doc.add_table(rows=len(selected), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.15, 1.15, 1.30, 1.15, 0.55, 0.60]
    for i, row_data in enumerate(selected):
        row = table.rows[i]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = value
            cell.width = Inches(widths[j])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_run_font(r, size=7.3, bold=(i == 0 or value == 'Infection source, n (%)'))
            if i == 0:
                set_cell_shading(cell, 'D9EAF7')
            elif value == 'Infection source, n (%)':
                set_cell_shading(cell, 'EDEDED')
    set_repeat_header_and_nosplit(table)
    note = (
        'Data are median [interquartile range] or n (%). P values are two-sided comparisons between '
        'survivors/censored patients and patients who died within 60 days. Continuous variables were '
        'compared using Wilcoxon rank-sum tests and categorical variables using Fisher’s exact tests. '
        'For infection source, the P value shown in the first category row is the global P value for the '
        'six-level variable; category-specific P values are not shown. P values were descriptive and were '
        'not used to select covariates for molecular Cox models. The complete baseline table is provided '
        'as Supplementary Table S9. BUN, blood urea nitrogen; INR, international normalised ratio; '
        'PF, PaO₂/FiO₂ ratio; SIC, sepsis-induced coagulopathy; SMD, standardised mean difference; '
        'SOFA, Sequential Organ Failure Assessment.'
    )
    add_text_paragraph(doc, note, size=7.5, space_before=3, space_after=0, highlight=highlighted)


def build_manuscript(output: Path, highlighted=False):
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], landscape=False, margins=(0.85, 0.85, 0.9, 0.9), line_numbers=True)
    add_title_page(doc, highlighted)
    add_heading(doc, 'Abstract', 1)
    add_text_paragraph(doc, ABSTRACT, highlight=highlighted, size=10.5)
    add_text_paragraph(doc, 'Keywords: ' + KEYWORDS, highlight=highlighted, size=10.5, italic=True, space_after=10)

    add_heading(doc, 'Introduction', 1)
    for i in intro_ids:
        add_text_paragraph(doc, mapped(source_text(i)), highlight=False)

    add_heading(doc, 'Results', 1)
    for heading, ids in result_blocks:
        add_heading(doc, heading, 2)
        for i in ids:
            add_text_paragraph(doc, mapped(source_text(i)), highlight=(highlighted and i in {72, 73}))

    add_heading(doc, 'Discussion', 1)
    for i in discussion_ids:
        if i == 135:
            continue
        add_text_paragraph(doc, mapped(source_text(i)), highlight=False)
    add_text_paragraph(doc, mapped(TRANSLATIONAL), highlight=highlighted)
    add_text_paragraph(doc, mapped(source_text(135)), highlight=highlighted)

    add_heading(doc, 'Methods', 1)
    for heading, ids in method_blocks:
        add_heading(doc, heading, 2)
        for i in ids:
            add_text_paragraph(doc, mapped(source_text(i)), highlight=(highlighted and i == 42))

    add_heading(doc, 'Ethics approval and informed consent', 2, highlight=highlighted)
    add_text_paragraph(doc, ETHICS, highlight=highlighted)
    add_heading(doc, 'Code availability', 2, highlight=highlighted)
    add_text_paragraph(doc, CODE_AVAILABILITY, highlight=highlighted)
    add_heading(doc, 'AI-assisted tools in manuscript and code preparation', 2, highlight=highlighted)
    add_text_paragraph(doc, AI_DISCLOSURE, highlight=highlighted)

    add_heading(doc, 'Data availability', 1, highlight=highlighted)
    add_text_paragraph(doc, DATA_AVAILABILITY, highlight=highlighted)

    add_heading(doc, 'References', 1)
    parsed_refs = {n: parse_reference(raw_refs[n]) for n in raw_refs}
    for old in citation_order:
        add_reference_paragraph(doc, CITATION_MAP[old], parsed_refs[old])

    add_heading(doc, 'Acknowledgements', 1)
    add_text_paragraph(doc, SRC[153])
    add_heading(doc, 'Funding', 1)
    add_text_paragraph(doc, SRC[149])
    add_heading(doc, 'Author contributions', 1)
    add_text_paragraph(doc, SRC[151])
    add_heading(doc, 'Competing interests', 1)
    add_text_paragraph(doc, 'The authors declare no competing interests.', highlight=highlighted)

    add_heading(doc, 'Figure legends', 1)
    for i in [78, 86, 90, 99]:
        legend = mapped(source_text(i))
        p = add_text_paragraph(doc, legend, size=9.5, space_after=8)
        if p.runs:
            # Bold the title sentence before the first full stop.
            full = p.text
            first, sep, rest = full.partition('. ')
            if sep:
                p.clear()
                r = p.add_run(first + '.')
                set_run_font(r, size=9.5, bold=True)
                r = p.add_run(' ' + rest)
                set_run_font(r, size=9.5)
                style_paragraph(p, size=9.5, space_after=8)

    add_main_table(doc, highlighted=highlighted)

    # Document metadata.
    doc.core_properties.title = TITLE
    doc.core_properties.subject = 'Scientific Reports submission manuscript'
    doc.core_properties.author = 'Hao Lyu and co-authors'
    doc.core_properties.last_modified_by = 'Hao Lyu and co-authors'
    doc.core_properties.keywords = KEYWORDS
    doc.save(output)


def add_full_table_to_doc(doc: Document, table_data: list[list[str]], title: str):
    table_data = [([row[0], '', '', '', '', ''] if row[0] == 'Infection source, n (%)' else row) for row in table_data]
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, landscape=True, margins=(0.34, 0.34, 0.28, 0.28), line_numbers=False)
    add_text_paragraph(doc, title, size=10.5, bold=True, space_after=4, keep_with_next=True)
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.0, 1.18, 1.28, 1.10, 0.56, 0.58]
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = value
            cell.width = Inches(widths[j])
            set_cell_margins(cell, 12, 28, 12, 28)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_run_font(r, size=6.3, bold=(i == 0 or value == 'Infection source, n (%)'))
            if i == 0:
                set_cell_shading(cell, 'D9EAF7')
            elif value == 'Infection source, n (%)':
                set_cell_shading(cell, 'EDEDED')
    set_repeat_header_and_nosplit(table)
    note = (
        'Data are median [interquartile range] or n (%). P values are two-sided comparisons between '
        'survivors/censored patients and patients who died within 60 days. Continuous variables were '
        'compared using Wilcoxon rank-sum tests and categorical variables using Fisher’s exact tests. '
        'For infection source, the P value shown in the first category row is the global P value for the '
        'six-level variable. P values were descriptive and were not used for covariate selection. '
        'aPTT, activated partial thromboplastin time; BMI, body mass index; BUN, blood urea nitrogen; '
        'Cl, chloride; COPD, chronic obstructive pulmonary disease; CRP, C-reactive protein; '
        'HRmax, maximum heart rate; INR, international normalised ratio; K, potassium; MAPmax, maximum '
        'mean arterial pressure; Na, sodium; PCT, procalcitonin; PF, PaO₂/FiO₂ ratio; RRmax, maximum '
        'respiratory rate; SBPmax, maximum systolic arterial pressure; SIC, sepsis-induced coagulopathy; '
        'SMD, standardised mean difference; SOFA, Sequential Organ Failure Assessment; Tmax, maximum '
        'temperature; WBC, white blood cell.'
    )
    add_text_paragraph(doc, note, size=6.4, space_before=2, space_after=0)


def build_supplementary(output: Path):
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], landscape=False, margins=(0.65, 0.65, 0.7, 0.7), line_numbers=False)
    add_text_paragraph(doc, 'Supplementary Information', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       size=14, bold=True, space_after=12)
    add_text_paragraph(doc, TITLE, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=10)
    add_text_paragraph(doc, AUTHORS, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=True, space_after=12)

    add_heading(doc, 'Supplementary Methods', 1)
    for heading_i, text_i in [(2, 3), (4, 5), (6, 7), (8, 9)]:
        add_heading(doc, SUPP_SRC[heading_i], 2)
        add_text_paragraph(doc, SUPP_SRC[text_i])

    add_heading(doc, 'Supplementary Figures', 1)
    add_text_paragraph(
        doc,
        'Supplementary Figs. S1–S9 are presented in numerical order. Supplementary Figs. S1–S8 '
        'provide prespecified Day-5 RNA-seq availability and weighting diagnostics, Hallmark robustness '
        'analyses, entry-boundary sensitivity and the descriptive Day-5 plasma protein availability audit. '
        'Supplementary Fig. S9 presents the complete displayed clinical univariable Cox contrasts.'
    )
    figure_root = ROOT / 'submission' / 'figures'
    images = [
        figure_root / 'Supplementary_Figure_A1_centre_positivity.png',
        figure_root / 'Supplementary_Figure_A2_probability_weight_distributions.png',
        figure_root / 'Supplementary_Figure_A3_pre_post_weight_SMD.png',
        figure_root / 'Supplementary_Figure_A4_all_Hallmark_unweighted_vs_IPW.png',
        figure_root / 'Supplementary_Figure_A5_core_pathway_scenario_heatmap.png',
        figure_root / 'Supplementary_Figure_A6_six_scenario_robustness_metrics.png',
        figure_root / 'Supplementary_Figure_A7_entry_boundary_sensitivity.png',
        figure_root / 'Supplementary_Figure_A8_D5_protein_availability.png',
        figure_root / 'Supplementary_Figure_A9_clinical_univariable_Cox.png',
    ]
    caption_indices = [15, 21, 33, 37, 45, 53, 68, 75, 88]
    widths = [6.7, 6.6, 5.4, 6.5, 6.5, 6.6, 6.0, 6.0, 5.2]
    for n, (image, cap_i, width) in enumerate(zip(images, caption_indices, widths), start=1):
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(image), width=Inches(width))
        cap = re.sub(r'Supplementary Figure A[1-9]', f'Supplementary Fig. S{n}', SUPP_SRC[cap_i])
        p = add_text_paragraph(doc, cap, size=9, space_before=4, space_after=4)
        first, sep, rest = p.text.partition('. ')
        if sep:
            p.clear()
            r = p.add_run(first + '.')
            set_run_font(r, size=9, bold=True)
            r = p.add_run(' ' + rest)
            set_run_font(r, size=9)
            style_paragraph(p, size=9, space_after=4)

    add_heading(doc, 'Supplementary Tables', 1)
    add_text_paragraph(doc, 'Supplementary Tables S1–S8 are supplied as separate machine-readable Excel workbooks. Supplementary Table S9 is included below and supplied as a separate Excel workbook.')
    index_rows = [
        ('Supplementary Table S1', 'Clinical variable definitions', 'Variable definitions, prespecified increments, transformations, infection-source harmonisation and missing-variable exclusions.'),
        ('Supplementary Table S2', 'Clinical univariable Cox models', 'Complete exploratory univariable Cox contrasts, proportional-hazards diagnostics and nonlinearity tests.'),
        ('Supplementary Table S3', 'RNA-seq gene-wise Cox and PH outputs', 'Complete gene-level Day-1, Day-3 and Day-5 centre-stratified Cox results and PH diagnostics.'),
        ('Supplementary Table S4', 'RNA-seq Hallmark enrichment', 'Complete Hallmark enrichment and leading-edge results for transcriptomic analyses.'),
        ('Supplementary Table S5', 'Protein-wise Cox and PH outputs', 'Complete protein-level Day-1, Day-3 and Day-5 centre-stratified Cox results and PH diagnostics.'),
        ('Supplementary Table S6', 'Protein Hallmark enrichment', 'Complete Hallmark enrichment and leading-edge results for proteomic analyses.'),
        ('Supplementary Table S7', 'Cross-omic models', 'Contemporaneous, forward, reverse-order and sensitivity cross-omic pathway models.'),
        ('Supplementary Table S8', 'Day-5 availability and IPW', 'Positivity, observation-model, weight, balance and weighted-enrichment diagnostics.'),
        ('Supplementary Table S9', 'Complete baseline characteristics', 'Complete Day-1 baseline characteristics according to 60-day outcome; this is the full source table from which main-text Table 1 was selected.'),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, val in enumerate(['Item', 'Title', 'Contents']):
        hdr[j].text = val
        set_cell_shading(hdr[j], 'D9EAF7')
    for row in index_rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            set_cell_margins(cell, 55, 65, 55, 65)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_run_font(r, size=8.5, bold=(i == 0))
    set_repeat_header_and_nosplit(t)

    source_doc = Document(SOURCE_MANUSCRIPT)
    full_table = [[cell.text for cell in row.cells] for row in source_doc.tables[0].rows]
    add_full_table_to_doc(doc, full_table, 'Supplementary Table S9. Complete baseline characteristics of the Day-1 SIC cohort according to 60-day outcome')
    doc.core_properties.title = 'Scientific Reports Supplementary Information'
    doc.core_properties.author = 'Hao Lyu and co-authors'
    doc.core_properties.last_modified_by = 'Hao Lyu and co-authors'
    doc.save(output)


def build_revision_report(output: Path):
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], margins=(0.8, 0.8, 0.8, 0.8), line_numbers=False)
    add_text_paragraph(doc, 'Scientific Reports transfer revision report', size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_text_paragraph(doc, TITLE, size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_text_paragraph(doc, 'Scientific freeze', size=12, bold=True, space_after=4)
    add_text_paragraph(doc, 'No cohort count, event count, effect estimate, confidence interval, P value, false-discovery rate, normalised enrichment score, pathway direction, estimand, model definition or scientific figure was changed. The conversion is editorial and journal-specific.')
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, v in enumerate(['Location', 'Classification', 'Change', 'Scientific boundary']):
        table.rows[0].cells[j].text = v
        set_cell_shading(table.rows[0].cells[j], 'D9EAF7')
    boundary = {
        'clarification': 'Clarifies the existing design or reporting without changing results.',
        'interpretive expansion': 'Adds bounded interpretation supported by frozen findings and existing citations.',
        'claim restriction': 'Reduces overinterpretation and explicitly excludes unsupported causal or clinical claims.',
        'journal compliance': 'Implements a Scientific Reports submission requirement.',
        'structural relocation': 'Moves existing text without changing its scientific meaning.',
        'cross-reference renumbering': 'Updates identifiers consistently without changing content.'
    }
    for loc, cls, change in CHANGE_LOG:
        cells = table.add_row().cells
        vals = [loc, cls, change, boundary[cls]]
        for j, v in enumerate(vals):
            cells[j].text = v
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, size=8.5, bold=(i == 0))
    set_repeat_header_and_nosplit(table)
    add_text_paragraph(doc, 'Highlighting convention', size=12, bold=True, space_before=10, space_after=4)
    add_text_paragraph(doc, 'The highlighted manuscript uses yellow highlighting for newly written or substantively rewritten scientific text, methodological-boundary clarifications, claim restrictions and new policy disclosures. Routine section moves, Nature-style reference conversion and systematic supplementary renumbering are documented here rather than highlighted throughout.')
    doc.save(output)


def build_cover_letter(output: Path):
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], margins=(0.62, 0.62, 0.62, 0.62), line_numbers=False)
    add_text_paragraph(doc, '25 July 2026', size=9.4, space_after=6)
    add_text_paragraph(doc, 'Editors\nScientific Reports', size=9.4, space_after=6)
    add_text_paragraph(doc, 'Dear Editors,', size=9.4, space_after=5)
    paragraphs = [
        f'On behalf of my co-authors, I am pleased to submit the Article entitled “{TITLE}” for consideration in Scientific Reports.',
        'This manuscript was previously submitted to the Journal of Intensive Care and is now being submitted to Scientific Reports following a Springer Nature journal-transfer recommendation. The earlier journal did not provide scientific peer-review comments requiring a response, and the present submission has been reformatted and clarified for Scientific Reports.',
        'The study addresses a methodological and biological problem relevant to a broad readership in sepsis, haemostasis, systems biology and prognostic research. Using a fixed Day-1 cohort of 504 adults with sepsis-induced coagulopathy, we applied separate risk-set-aware landmark survival analyses to whole-blood transcriptomic and plasma proteomic measurements obtained on Days 1, 3 and 5. The design distinguishes later-landmark prognosis from within-patient molecular trajectories, explicitly separates structural non-availability caused by death from sample observation, and evaluates the robustness of Day-5 RNA-seq findings through positivity-supported inverse-probability weighting.',
        'The principal contribution is a technically validated, reproducible description of stage- and compartment-dependent molecular associations with 60-day mortality. Recurrent whole-blood heme, hypoxia and inflammatory enrichment coexisted with broader negative enrichment of cellular-maintenance programmes at Day 5; later plasma protein rankings were dominated by extracellular-matrix and tissue-remodelling signals; and interferon pathways showed selective contemporaneous and forward cross-omic associations. The conclusions are deliberately bounded: the observational models do not establish causal transfer, within-patient trajectories, incremental clinical prediction or therapeutic benefit.',
        f'Participant-level clinical, transcriptomic and proteomic data remain controlled access under OMIX011182 and cannot be redistributed by the authors. Version-controlled code, aggregate source data, numerical-truth interfaces and quality-assurance outputs are archived in the Scientific Reports-specific repository release at {RELEASE_URL}.',
        'The work is original, has not been published, and is not under consideration by another journal. All authors have reviewed and approved the submitted version and agree to be accountable for their contributions. We have no preferred reviewers and request no reviewer exclusions. We have had no prior discussions with a Scientific Reports Editorial Board Member regarding this work.',
        'Thank you for considering this manuscript. We believe its multicentre longitudinal design, explicit landmark estimands, multi-omic integration and transparent reproducibility framework are well aligned with Scientific Reports’ emphasis on technical robustness and scientific validity.'
    ]
    for text in paragraphs:
        add_text_paragraph(doc, text, size=9.4, space_after=5)
    add_text_paragraph(doc, 'Sincerely,', size=9.4, space_before=5, space_after=6)
    add_text_paragraph(doc, 'Dong Zhang, MD, PhD\nCorresponding author\nDepartment of Critical Care Medicine\nThe First Hospital of Jilin University\nNo. 1 Xinmin Street, Changchun, Jilin 130021, China\nEmail: zhangdong@jlu.edu.cn\nOn behalf of all authors', size=9.4, space_after=0)
    doc.core_properties.title = 'Cover letter to Scientific Reports'
    doc.core_properties.author = 'Dong Zhang'
    doc.core_properties.last_modified_by = 'Dong Zhang'
    doc.save(output)


def write_mapping():
    path = WORK / 'citation_renumbering.tsv'
    with path.open('w', encoding='utf-8', newline='') as f:
        w = csv.writer(f, delimiter='\t')
        w.writerow(['old_reference_number', 'new_reference_number'])
        for old in citation_order:
            w.writerow([old, CITATION_MAP[old]])


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Build Scientific Reports submission documents from the frozen JIC interface.')
    parser.add_argument('--mode', choices=['manuscript', 'supplementary', 'cover', 'all'], default='manuscript')
    parser.add_argument('--cover-out', type=Path, default=None)
    args = parser.parse_args()
    if args.mode in {'manuscript', 'all'}:
        build_manuscript(OUT / 'Scientific_Reports_manuscript_clean.docx', highlighted=False)
        build_manuscript(OUT / 'Scientific_Reports_manuscript_highlighted.docx', highlighted=True)
        build_revision_report(OUT / 'Scientific_Reports_revision_report.docx')
        write_mapping()
    if args.mode in {'supplementary', 'all'}:
        build_supplementary(OUT / 'Scientific_Reports_Supplementary_Information.docx')
    if args.mode in {'cover', 'all'}:
        if args.cover_out is None:
            raise SystemExit('--cover-out is required for cover generation because the cover letter is private and must not be written into the repository.')
        args.cover_out.parent.mkdir(parents=True, exist_ok=True)
        build_cover_letter(args.cover_out)
    print('Scientific Reports build completed:', args.mode)



if __name__ == '__main__':
    main()
