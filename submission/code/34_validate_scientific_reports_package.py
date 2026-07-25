"""Validate the frozen Scientific Reports submission package without refitting analyses."""
from __future__ import annotations

import argparse
import csv
import json
import re
import zipfile
from pathlib import Path

from docx import Document

TITLE = "Landmark-specific transcriptomic and proteomic associations with 60-day mortality in Day-1-defined sepsis-induced coagulopathy: a multicentre longitudinal cohort study"


def words(text: str) -> int:
    return len(re.findall(r"\b[\wα-ωΑ-Ωρ]+(?:[-–][\wα-ωΑ-Ωρ]+)*\b", text, flags=re.UNICODE))


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(parts)


def ooxml_counts(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml = "".join(archive.read(name).decode("utf-8", "ignore") for name in names if name.endswith(".xml"))
    return {
        "comments": sum(name.endswith("comments.xml") for name in names),
        "tracked": len(re.findall(r"<w:(?:ins|del)\b", xml)),
        "hidden": len(re.findall(r"<w:vanish\b", xml)),
        "yellow": len(re.findall(r'<w:highlight[^>]+w:val="yellow"', xml)),
    }


def section_text(document: Document, start: str, end: str) -> str:
    paragraphs = document.paragraphs
    positions = {p.text: i for i, p in enumerate(paragraphs) if p.style and p.style.name == "Heading 1"}
    return "\n".join(p.text for p in paragraphs[positions[start] + 1 : positions[end]] if p.text.strip())


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--repo-root", type=Path, default=Path("."))
    parser.add_argument("--private-cover-letter", type=Path, required=True)
    parser.add_argument("--out", type=Path, default=Path("submission/qa/scientific_reports_package_validation.json"))
    args = parser.parse_args()
    repo = args.repo_root.resolve()
    package = repo / "submission/manuscript_files/scientific_reports"
    manuscript_path = package / "Scientific_Reports_manuscript_clean.docx"
    highlighted_path = package / "Scientific_Reports_manuscript_highlighted.docx"
    supplement_path = package / "Scientific_Reports_Supplementary_Information.docx"
    manuscript = Document(manuscript_path)
    paragraphs = manuscript.paragraphs
    h1 = [p.text for p in paragraphs if p.style and p.style.name == "Heading 1"]
    expected_h1 = ["Abstract", "Introduction", "Results", "Discussion", "Methods", "Data availability", "References", "Acknowledgements", "Funding", "Author contributions", "Competing interests", "Figure legends"]
    positions = {p.text: i for i, p in enumerate(paragraphs) if p.style and p.style.name == "Heading 1"}
    abstract = paragraphs[positions["Abstract"] + 1].text
    keywords = paragraphs[positions["Abstract"] + 2].text.replace("Keywords:", "").strip().split(";")
    ird = sum(words(section_text(manuscript, a, b)) for a, b in (("Introduction", "Results"), ("Results", "Discussion"), ("Discussion", "Methods")))
    methods_h2 = [p.text for p in paragraphs[positions["Methods"] + 1 : positions["Data availability"]] if p.style and p.style.name == "Heading 2"]
    manuscript_text = document_text(manuscript_path)
    supplement_text = document_text(supplement_path)
    clean = ooxml_counts(manuscript_path)
    highlighted = ooxml_counts(highlighted_path)
    numeric_audit = list(csv.DictReader((repo / "submission/qa/scientific_reports_numeric_audit.tsv").open(encoding="utf-8"), delimiter="\t"))
    cross_audit = list(csv.DictReader((repo / "submission/qa/scientific_reports_cross_reference_audit.tsv").open(encoding="utf-8"), delimiter="\t"))
    forbidden = ("D:/", "C:\\Users\\", "SampleName", "PatientID", "participant_id", "password", "token=")
    public_text = manuscript_text + "\n" + supplement_text
    cover_text = document_text(args.private_cover_letter)
    report = {
        "title": paragraphs[0].text,
        "abstract_words": words(abstract),
        "keyword_count": len([x for x in keywords if x.strip()]),
        "intro_results_discussion_words": ird,
        "section_order_pass": h1 == expected_h1,
        "data_availability_before_references": positions["Data availability"] < positions["References"],
        "code_availability_in_methods": "Code availability" in methods_h2,
        "ai_is_final_methods_subsection": bool(methods_h2) and methods_h2[-1] == "AI-assisted tools in manuscript and code preparation",
        "numeric_mismatches": sum(str(row["match"]).lower() not in {"true", "1"} for row in numeric_audit),
        "broken_cross_references": sum(row["status"] != "PASS" for row in cross_audit),
        "residual_A_labels": len(re.findall(r"Supplementary (?:Figure|Fig\.) A[1-9]", public_text)),
        "privacy_findings": sum(term in public_text for term in forbidden),
        "clean_comments": clean["comments"],
        "clean_tracked_changes": clean["tracked"],
        "clean_hidden_text": clean["hidden"],
        "highlighted_yellow_runs": highlighted["yellow"],
        "supplementary_information_bytes": (package / "Scientific_Reports_Supplementary_Information.pdf").stat().st_size,
        "cover_letter_private": not str(args.private_cover_letter.resolve()).startswith(str(repo)),
        "cover_transfer_statement": "Springer Nature journal-transfer recommendation" in cover_text,
        "cover_reviewer_statement": "no preferred reviewers and request no reviewer exclusions" in cover_text,
    }
    report["all_pass"] = all((
        report["title"] == TITLE,
        report["abstract_words"] <= 200,
        report["keyword_count"] == 6,
        report["intro_results_discussion_words"] <= 4500,
        report["section_order_pass"],
        report["data_availability_before_references"],
        report["code_availability_in_methods"],
        report["ai_is_final_methods_subsection"],
        report["numeric_mismatches"] == 0,
        report["broken_cross_references"] == 0,
        report["residual_A_labels"] == 0,
        report["privacy_findings"] == 0,
        report["clean_comments"] == 0,
        report["clean_tracked_changes"] == 0,
        report["clean_hidden_text"] == 0,
        report["highlighted_yellow_runs"] > 0,
        report["supplementary_information_bytes"] < 50 * 1024 * 1024,
        report["cover_letter_private"],
        report["cover_transfer_statement"],
        report["cover_reviewer_statement"],
    ))
    output = repo / args.out
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["all_pass"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
