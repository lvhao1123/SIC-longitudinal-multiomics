"""Validate production-only manuscript revisions without refitting analyses."""

from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def count_comments(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        if "word/comments.xml" not in archive.namelist():
            return 0
        xml = archive.read("word/comments.xml").decode("utf-8")
    return xml.count("<w:comment ")


def embedded_image_matches(docx_path: Path, image_path: Path) -> bool:
    expected = hashlib.sha256(image_path.read_bytes()).hexdigest()
    with zipfile.ZipFile(docx_path) as archive:
        observed = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
    return expected in observed


def main_figure_blocks_are_bound(doc) -> bool:
    paragraphs = doc.paragraphs
    titles = {
        1: "Study design, longitudinal risk sets and Day-5 availability estimand",
        2: "Time-specific whole-blood transcriptomic prognostic programmes",
        3: "Time-specific plasma protein prognostic programmes",
        4: "Pathway-selective contemporaneous and forward cross-omic associations",
    }
    for number, title in titles.items():
        if any(p.text == f"Figure {number}" or p.text == f"Figure {number} | {title}" for p in paragraphs):
            return False
        indices = [
            i for i, p in enumerate(paragraphs)
            if p.text.startswith(f"Figure {number} | {title}.")
        ]
        if len(indices) != 1:
            return False
        i = indices[0]
        if i == 0:
            return False
        drawing, legend = paragraphs[i - 1], paragraphs[i]
        if not drawing._p.xpath(".//w:drawing") or not drawing.paragraph_format.keep_with_next:
            return False
        if not drawing.paragraph_format.page_break_before:
            return False
        if not legend.text.startswith(f"Figure {number} | {title}."):
            return False
        if legend.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            return False
        if abs(float(legend.paragraph_format.line_spacing) - 1.05) > 1e-6:
            return False
        if any(run.font.size is None or abs(run.font.size.pt - 9) > 1e-6 for run in legend.runs if run.text):
            return False
    return True


def validate(docx_path: Path, annotated: bool, main_figures: dict[int, Path], a7: Path, a9: Path):
    doc = Document(docx_path)
    body = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    table1_labels = (
        [row.cells[0].text for row in doc.tables[0].rows]
        if len(doc.tables) == 1 else []
    )
    checks = {
        "new_repository_url": "https://github.com/lvhao1123/SIC-longitudinal-multiomics" in body,
        "obsolete_repository_absent": "lvhao1123/SIC-research" not in body,
        "obsolete_tag_absent": "analysis-freeze-v1.0-2026-07-13" not in body,
        "controlled_data_statement": "formal CMAISE/OMIX application process" in body,
        "s1_to_s8_cited": all(f"Supplementary Table S{i}" in body for i in range(1, 9)),
        "a1_to_a9_cited": all(f"Supplementary Figure A{i}" in body for i in range(1, 10)),
        "a9_caption_present": "Exploratory clinical univariable Cox associations" in body,
        "a9_image_inserted": len(doc.inline_shapes) == 13,
        "a7_margin_corrected_image_inserted": embedded_image_matches(docx_path, a7),
        "a9_margin_corrected_image_inserted": embedded_image_matches(docx_path, a9),
        "no_time_updated_exposure_phrase": "time-updated exposures" not in body,
        "no_single_time_dependent_claim": "time-dependent covariate analysis" not in body,
        "four_main_figures_only": "Figure 5" not in body,
        "ph_excluded_from_table1": "Arterial pH" not in table_text,
        "calcium_excluded_from_table1": "Calcium" not in table_text,
        "table1_requested_dimensions": len(doc.tables) == 1 and len(doc.tables[0].rows) == 44 and len(doc.tables[0].columns) == 6,
        "table1_standard_abbreviations": all(
            label in table1_labels
            for label in [
                "BMI, kg/m²", "HRmax, beats/min", "MAPmax, mmHg", "SBPmax, mmHg",
                "RRmax, breaths/min", "Tmax, °C", "K, mmol/L", "Na, mmol/L", "Cl, mmol/L",
                "BUN, mg/dL", "CRP, mg/L", "PCT, ng/mL",
                "WBC count, ×10⁹/L", "INR", "aPTT, s",
                "PaO₂/FiO₂ ratio, mmHg", "COPD",
            ]
        ),
        "table1_section_rows_removed": "Continuous variables" not in table1_labels and "Binary variables" not in table1_labels,
        "table1_numeric_cells_centred": len(doc.tables) == 1 and all(
            p.alignment == WD_ALIGN_PARAGRAPH.CENTER
            for row in doc.tables[0].rows
            for cell in row.cells[1:]
            for p in cell.paragraphs
        ),
        "table1_uniform_font_size": len(doc.tables) == 1 and all(
            run.font.size is not None and abs(run.font.size.pt - 8) < 1e-6
            for row in doc.tables[0].rows
            for cell in row.cells
            for p in cell.paragraphs
            for run in p.runs if run.text
        ),
        "table1_model_coding_labels_removed": "1 vs 0" not in table_text,
        "table1_abbreviation_note_present": all(
            term in body
            for term in [
                "BMI, body mass index", "HRmax, maximum heart rate",
                "MAPmax, maximum mean arterial pressure", "RRmax, maximum respiratory rate",
                "SBPmax, maximum systolic arterial pressure", "Tmax, maximum temperature",
                "PCT, procalcitonin", "WBC, white blood cell",
            ]
        ),
        "r_version_reported": "R version 4.4.2" in body,
        "pathway_coverage_wording": (
            "All 15 pathways met the prespecified coverage requirement" in body
            and "exceeding the minimum of 10 detected members" not in body
        ),
        "main_figure_blocks_bound": main_figure_blocks_are_bound(doc),
        "main_figures_inserted": all(
            embedded_image_matches(docx_path, path) for path in main_figures.values()
        ),
        "risk_set_numbers_present": all(x in body for x in ["504", "421", "420", "321", "320", "491", "487", "167"]),
        "zero_centre_patient_distinction": "Four patients from three zero-observation centres" in body,
        "a9_page_break_before": any(
            p.text.startswith("Supplementary Figure A9")
            and p.paragraph_format.page_break_before
            for p in doc.paragraphs
        ),
        "comments_expected_state": count_comments(docx_path) >= 15 if annotated else count_comments(docx_path) == 0,
    }
    return checks


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--clean", type=Path, required=True)
    parser.add_argument("--annotated", type=Path, required=True)
    parser.add_argument("--fig1", type=Path, required=True)
    parser.add_argument("--fig2", type=Path, required=True)
    parser.add_argument("--fig3", type=Path, required=True)
    parser.add_argument("--fig4", type=Path, required=True)
    parser.add_argument("--a7", type=Path, required=True)
    parser.add_argument("--a9", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    main_figures = {1: args.fig1, 2: args.fig2, 3: args.fig3, 4: args.fig4}
    clean = validate(args.clean, annotated=False, main_figures=main_figures, a7=args.a7, a9=args.a9)
    annotated = validate(args.annotated, annotated=True, main_figures=main_figures, a7=args.a7, a9=args.a9)
    report = {
        "clean": clean,
        "annotated": annotated,
        "all_pass": all(clean.values()) and all(annotated.values()),
    }
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["all_pass"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
